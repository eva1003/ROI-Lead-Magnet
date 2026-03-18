#!/usr/bin/env python3
"""
Erstellt CSRD_Report_v7.4.docx aus v7.3.
Ändert ausschließlich den Abschnitt "Allgemeine Informationen (ESRS 2)":
  - Bullet-Listen → Fließtext
  - Nummerierte Listenabsätze → integrierten Fließtext
  - Kleinere Duplikate / GOV-3 Zusammenfassung
  - Lesefluss-Verbesserungen
"""
import shutil
from pathlib import Path
from lxml import etree

SRC = Path("Fröbel 2024/CSRD_Report_v7.3.docx")
DST = Path("Fröbel 2024/CSRD_Report_v7.4.docx")

shutil.copy(SRC, DST)

# ── XML-Helfer ─────────────────────────────────────────────────────────────
from docx import Document
from docx.oxml import OxmlElement

doc = Document(DST)
body = doc.element.body
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def gtext(el):
    return "".join(t.text or "" for t in el.findall(f".//{{{W}}}t"))


def para_style(el):
    ps = el.find(f".//{{{W}}}pStyle")
    return ps.get(f"{{{W}}}val", "") if ps is not None else ""


def clear_runs(p_el):
    """Alle w:r und w:hyperlink entfernen."""
    for child in list(p_el):
        tag = child.tag.split("}")[-1]
        if tag in ("r", "hyperlink", "ins", "del"):
            p_el.remove(child)


def set_text(p_el, text):
    """Paragraph-Text ersetzen (pPr bleibt erhalten)."""
    clear_runs(p_el)
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p_el.append(r)


def set_style_normal(p_el):
    """pStyle auf Normal setzen (Bullet-Formatierung entfernen)."""
    pPr = p_el.find(f"{{{W}}}pPr")
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    # pStyle
    pStyle = pPr.find(f"{{{W}}}pStyle")
    if pStyle is not None:
        pStyle.set(f"{{{W}}}val", "Normal")
    # numPr entfernen (Aufzählungsformatierung)
    numPr = pPr.find(f"{{{W}}}numPr")
    if numPr is not None:
        pPr.remove(numPr)
    # ind entfernen falls vorhanden (Einrückung durch Bullet)
    ind = pPr.find(f"{{{W}}}ind")
    if ind is not None:
        pPr.remove(ind)


def remove_elements(indices):
    """Body-Elemente nach Index entfernen (absteigend, damit Index stabil)."""
    children = list(body)
    for i in sorted(set(indices), reverse=True):
        body.remove(children[i])


# ── Aktuellen Elementindex-Snapshot ────────────────────────────────────────
def snap():
    return list(body)


# ==========================================================================
# 1. GOV-2 Bullets (Elemente 60-71) → Fließtext, Intro anpassen
# ==========================================================================
children = snap()

GOV2_INTRO_IDX = 59
GOV2_BULLET_INDICES = list(range(60, 72))

new_gov2_text = (
    "Im Berichtszeitraum befassten sich die Verwaltungs-, Leitungs- und/oder "
    "Aufsichtsorgane mit einer Vielzahl wesentlicher Auswirkungen, Risiken und Chancen. "
    "Im Umweltbereich standen insbesondere der Energieverbrauch und die "
    "Treibhausgasemissionen aus dem Betrieb von Gebäuden, konkrete Klimaschutzmaßnahmen "
    "zur Reduktion klimabezogener Auswirkungen sowie klimabezogene Risiken für den Betrieb "
    "– namentlich im Zusammenhang mit Hitze und Extremwetterereignissen – im Mittelpunkt. "
    "Im sozialen Bereich wurden Arbeitsbedingungen und die Beschäftigungssituation der "
    "Mitarbeitenden, Fachkräftesicherung und Personalgewinnung, die Qualität der "
    "frühkindlichen Bildung, Bildung für nachhaltige Entwicklung (BNE) sowie die "
    "gesellschaftliche Wirkung der pädagogischen Arbeit erörtert. Auf strategischer Ebene "
    "befassten sich die Organe zudem mit der Weiterentwicklung der Nachhaltigkeitsstrategie, "
    "der Durchführung der doppelten Wesentlichkeitsanalyse, nachhaltigkeitsbezogenen Risiken "
    "und Chancen sowie den regulatorischen Anforderungen aus der CSRD, den ESRS und der "
    "Nachhaltigkeitsberichterstattung."
)

set_text(children[GOV2_INTRO_IDX], new_gov2_text)
set_style_normal(children[GOV2_INTRO_IDX])
remove_elements(GOV2_BULLET_INDICES)

# ==========================================================================
# 2. GOV-3 (Elemente 74-77 in Original → nach Entfernen von 12 Bullets: -12)
#    Jetzt indizieren wir neu
# ==========================================================================
children = snap()

# Nach Entfernen von 12 Bullet-Elementen: urspr. Idx 74 → 74-12=62
# Überprüfen:
for i, el in enumerate(children):
    t = gtext(el)
    if "keine nachhaltigkeitsbezogenen Anreizsysteme" in t:
        GOV3_START = i
        break

# 4 Normal-Absätze ab GOV3_START zusammenfassen auf 2
c = children
g0 = gtext(c[GOV3_START])      # "...keine Anreizsysteme...nicht an Kriterien gekoppelt"
g1 = gtext(c[GOV3_START + 1])  # "Keine Bewertung nach nachhaltigkeitsbezogenen Zielen"
g2 = gtext(c[GOV3_START + 2])  # "Keine KPIs als Leistungsrichtwerte"
g3 = gtext(c[GOV3_START + 3])  # "Da keine Systeme → keine Genehmigung"

new_gov3_p1 = (
    "Für die Mitglieder der Verwaltungs-, Leitungs- und Aufsichtsorgane von Fröbel "
    "bestehen derzeit keine nachhaltigkeitsbezogenen Anreizsysteme. Die Vergütung der "
    "Organmitglieder ist nicht an nachhaltigkeitsbezogene Kriterien oder Zielsetzungen "
    "gekoppelt; nachhaltigkeitsbezogene Leistungskennzahlen werden weder als "
    "Leistungsrichtwerte herangezogen noch in die Vergütungspolitik einbezogen."
)
new_gov3_p2 = (
    "Eine Bewertung der Leistung von Organmitgliedern anhand spezifischer "
    "nachhaltigkeitsbezogener Ziele oder Auswirkungen erfolgt derzeit nicht. "
    "Da keine entsprechenden Anreizsysteme bestehen, entfällt auch eine formale "
    "Genehmigung oder Aktualisierung solcher Bedingungen auf Leitungsebene."
)

set_text(c[GOV3_START], new_gov3_p1)
set_text(c[GOV3_START + 1], new_gov3_p2)
remove_elements([GOV3_START + 2, GOV3_START + 3])

# ==========================================================================
# 3. GOV-4 nummerierte Schritte: Trennzeichen nach Titel ergänzen
#    Muster: "N. Titel Inhalt" → "N. Titel: Inhalt"
# ==========================================================================
children = snap()

gov4_step_keywords = [
    "Einbettung in Governance",
    "Ermittlung und Bewertung von Auswirkungen",
    "Präventions- und Minderungsmaßnahmen",
    "Überwachung der Wirksamkeit",
    "Kommunikation",
    "Abhilfemechanismen",
]

for i, el in enumerate(children):
    t = gtext(el)
    for kw in gov4_step_keywords:
        if t.startswith(("1. " + kw[:6], "2. " + kw[:6], "3. " + kw[:6],
                         "4. " + kw[:6], "5. " + kw[:6], "6. " + kw[:6])):
            # Titel-Ende finden und Doppelpunkt einfügen
            # Format: "N. TitelWörter Inhalt" → Titel endet wo Inhalt beginnt
            # Einfache Heuristik: nach dem zweiten Großbuchstaben-Wort
            idx_kw = t.find(kw)
            if idx_kw == -1:
                break
            # Finde Ende des Titels (zweites Großwort nach Keyword → vor "Die/Das/Ein")
            rest = t[idx_kw + len(kw):]
            # rest beginnt entweder mit Leerzeichen + Großbuchstabe (Inhalt) oder Ende
            import re
            m = re.search(r'\s+([A-ZÜÄÖ])', rest)
            if m:
                split_pos = idx_kw + len(kw) + m.start()
                new_t = t[:split_pos] + ": " + t[split_pos:].lstrip()
                set_text(el, new_t)
            break

# ==========================================================================
# 4. IRO-1 nummerierte Prozessschritte (5 Schritte) → Fließtext
# ==========================================================================
children = snap()

# Finde Paragraph "Die Durchführung erfolgte in fünf aufeinanderfolgenden Schritten:"
iro1_intro_idx = None
for i, el in enumerate(children):
    if "fünf aufeinanderfolgenden Schritten" in gtext(el):
        iro1_intro_idx = i
        break

if iro1_intro_idx is not None:
    # Die nächsten 5 Elemente sind die Schritte (Normal-Absätze)
    step_texts = [gtext(children[iro1_intro_idx + j]) for j in range(1, 6)]

    new_iro1_steps = (
        "Die Durchführung erfolgte in fünf aufeinanderfolgenden Schritten: erstens der "
        "Festlegung des Konsolidierungskreises und Abbildung der relevanten "
        "Wertschöpfungsketten (vgl. ESRS 1, Rz. 63 ff.); zweitens der Identifikation und "
        "Bewertung relevanter Stakeholder (vgl. ESRS 2, Rz. 43 ff.); drittens der "
        "Identifikation von Auswirkungen, Risiken und Chancen anhand der Aktivitäten in der "
        "Wertschöpfungskette (Top-down und Bottom-up, vgl. EFRAG IG 1, Abschnitt 3,2, "
        "Rz. 73 ff.); viertens der Bewertung der identifizierten IROs nach Relevanz, "
        "Schweregrad, Umfang und Wahrscheinlichkeit (vgl. ESRS 1, Abschnitt 3,4, "
        "Rz. 84 ff.); sowie fünftens der Ableitung der Berichtspflichten und Überführung "
        "in die Nachhaltigkeitserklärung (vgl. ESRS 2, Abschnitt 4,1, Rz. 54 ff.)."
    )

    set_text(children[iro1_intro_idx], new_iro1_steps)
    remove_elements(list(range(iro1_intro_idx + 1, iro1_intro_idx + 6)))

# ==========================================================================
# 5. IRO-1 Überblick Bullets → Fließtext
# ==========================================================================
children = snap()

for i, el in enumerate(children):
    if gtext(el).strip() == "Das Vorgehen im Überblick:":
        overview_intro_idx = i
        break

# Intro-Element (i) + 4 Bullets (i+1..i+4) → 1 Absatz
bullets_ov = [gtext(children[overview_intro_idx + j]) for j in range(1, 5)]
new_overview = (
    "Im Einzelnen umfasste das Vorgehen die Identifikation der Auswirkungen entlang der "
    "Wertschöpfungskette (Scope 1–3), einschließlich Beschaffung von Waren, Lebensmitteln "
    "und Energie, Mobilität der Mitarbeitenden und Familien sowie Abfall und Recycling; die "
    "Bewertung der Auswirkungen anhand von Schweregrad, Umfang und Wahrscheinlichkeit "
    "(vgl. ESRS 1, Rz. 85 ff.); die Priorisierung nach relativer Bedeutung für Mensch und "
    "Umwelt (vgl. ESRS 1, Rz. 87 ff.); sowie das fortlaufende Monitoring durch "
    "regelmäßige Überprüfung und Dokumentation im IKS und Integration in das allgemeine "
    "Nachhaltigkeitsmanagement."
)
set_text(children[overview_intro_idx], new_overview)
set_style_normal(children[overview_intro_idx])
remove_elements(list(range(overview_intro_idx + 1, overview_intro_idx + 5)))

# ==========================================================================
# 6. IRO-1 Scoring-Bullets → in vorhergehenden Absatz integrieren
# ==========================================================================
children = snap()

scoring_intro_idx = None
for i, el in enumerate(children):
    if "mehrdimensionales Scoring-System auf einer 5-Punkte-Skala eingesetzt:" in gtext(el):
        scoring_intro_idx = i
        break

if scoring_intro_idx is not None:
    bullets_sc = [gtext(children[scoring_intro_idx + j]) for j in range(1, 5)]
    base_text = gtext(children[scoring_intro_idx])
    # Doppelpunkt am Ende entfernen, Bullets anhängen
    base_text = base_text.rstrip(":").rstrip()
    new_scoring = (
        base_text + ", das folgende Dimensionen umfasste: "
        "Eintrittswahrscheinlichkeit (von sehr unwahrscheinlich bis sicherer Eintritt), "
        "Ausmaß positiver und negativer Auswirkungen (von minimal bis sehr hoch), "
        "Umfang der Auswirkungen (von minimal bis maximal/weitreichend) sowie – "
        "ausschließlich für negative Auswirkungen – Wiederherstellbarkeit "
        "(von sehr einfach bis unmöglich)."
    )
    set_text(children[scoring_intro_idx], new_scoring)
    remove_elements(list(range(scoring_intro_idx + 1, scoring_intro_idx + 5)))

# ==========================================================================
# 7. IRO-1 Finanzielle Prozessschritte Bullets → in Intro-Absatz integrieren
# ==========================================================================
children = snap()

fin_intro_idx = None
for i, el in enumerate(children):
    t = gtext(el)
    if "Die Analyse folgte einem mehrstufigen Vorgehen:" in t:
        fin_intro_idx = i
        break

if fin_intro_idx is not None:
    bullets_fin = [gtext(children[fin_intro_idx + j]) for j in range(1, 5)]
    base_text = gtext(children[fin_intro_idx]).rstrip(":").rstrip()
    new_fin = (
        base_text + ": Zunächst wurden potenzielle finanzielle Risiken und Chancen entlang "
        "der Wertschöpfungskette erfasst, insbesondere in Bezug auf Personalgewinnung und "
        "-bindung, Energieversorgung, Lebensmittelversorgung sowie externe Regulierungen. "
        "Daran schloss sich die Bewertung nach Eintrittswahrscheinlichkeit, Ausmaß und "
        "zeitlichem Horizont (kurz-, mittel- und langfristig) an, gefolgt von der "
        "Priorisierung durch Abgleich der finanziellen Tragweite mit den im zentralen "
        "Risikomanagementsystem erfassten Unternehmensrisiken. Die laufende Überwachung "
        "erfolgt durch regelmäßige Aktualisierung im Rahmen des IKS und Integration in "
        "das Risikomanagement."
    )
    set_text(children[fin_intro_idx], new_fin)
    remove_elements(list(range(fin_intro_idx + 1, fin_intro_idx + 5)))

# ==========================================================================
# 8. IRO-1 Risikofelder Bullets (5 Stück) → Fließtext in Intro-Absatz
# ==========================================================================
children = snap()

risk_intro_idx = None
for i, el in enumerate(children):
    t = gtext(el)
    if "Im Einzelnen wurden folgende Felder betrachtet:" in t:
        risk_intro_idx = i
        break

if risk_intro_idx is not None:
    # 5 Bullets folgen
    new_risk = (
        "Das Verfahren berücksichtigt sowohl externe Abhängigkeiten als auch die sich "
        "daraus ergebenden finanziellen Risiken und Chancen (vgl.\u202fESRS 1, Rz.\u202f78–79). "
        "Im Bereich Arbeitsmarkt und Demografie besteht das Risiko des Fachkräftemangels "
        "mit finanziellen Auswirkungen auf Betriebssicherheit und Auslastung; als Chance "
        "eröffnet sich die Positionierung als attraktiver Arbeitgeber. Hinsichtlich der "
        "regulatorischen Rahmenbedingungen drohen steigende Kosten durch neue Vorgaben in "
        "der frühkindlichen Bildung oder arbeitsrechtliche Veränderungen, während staatliche "
        "Förderprogramme eine Chance darstellen. Bei den Energiepreisen besteht das Risiko "
        "deutlicher Kostensteigerungen für Strom und Wärme; dem stehen Chancen durch "
        "Investitionen in Energieeffizienz und langfristige Verträge mit erneuerbaren Energien "
        "gegenüber. Im Bereich Lebensmittelversorgung bestehen Risiken durch höhere "
        "Einkaufspreise und Lieferengpässe, während nachhaltige Beschaffungsstrategien und "
        "regionale Lieferketten Chancen eröffnen. Schließlich birgt der gesellschaftliche "
        "Trend rückläufiger Geburtenzahlen ein Nachfragerisiko, dem die steigende "
        "gesellschaftliche Wertschätzung früher Bildung als Chance gegenübersteht."
    )
    set_text(children[risk_intro_idx], new_risk)
    remove_elements(list(range(risk_intro_idx + 1, risk_intro_idx + 6)))

# ==========================================================================
# 9. IRO-1 Interne Datenquellen Bullets → Fließtext
# ==========================================================================
children = snap()

datasrc_idx = None
for i, el in enumerate(children):
    if gtext(el).strip() == "Als interne Datenquellen wurden herangezogen:":
        datasrc_idx = i
        break

if datasrc_idx is not None:
    new_datasrc = (
        "Als interne Datenquellen wurden herangezogen: die Treibhausgasbilanz (Scope 1–3) "
        "einschließlich Energieverbrauch, Lebensmittel und Mobilität, Personalstatistiken "
        "(z.\u202fB. Mitarbeiterzahlen, Fluktuation, Fachkräftebedarf), Nutzungs- und "
        "Belegungszahlen der Einrichtungen, Ergebnisse aus Mitarbeiter- und "
        "Familienbefragungen sowie die Dokumentation im internen Kontrollsystem (IKS)."
    )
    set_text(children[datasrc_idx], new_datasrc)
    remove_elements(list(range(datasrc_idx + 1, datasrc_idx + 6)))

# ==========================================================================
# 10. SBM-3: Informelle Formulierung verbessern
# ==========================================================================
children = snap()

for i, el in enumerate(children):
    t = gtext(el)
    if "Da das der erste CSRD-Report ist" in t:
        new_t = (
            "Da es sich um den ersten CSRD-Bericht von FRÖBEL e.V. handelt, entfällt ein "
            "Vergleich mit dem vorangegangenen Berichtszeitraum."
        )
        set_text(el, new_t)
        break

# ==========================================================================
# Speichern
# ==========================================================================
doc.save(DST)
print(f"✓ Gespeichert: {DST}")

# Schnellprüfung
doc2 = Document(DST)
bullets_remaining = sum(1 for p in doc2.paragraphs if p.style.name == "List Bullet")
print(f"  Verbleibende Bullet-Absätze im Dokument: {bullets_remaining}")
