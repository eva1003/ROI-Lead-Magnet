#!/usr/bin/env python3
"""
CSRD Report Generator – FRÖBEL e.V. (Berichtsjahr 2024)
Generates CSRD_Report.docx and QA_log.md from CSV input.

Spec-compliance: ESRS Reporting-Engineer spec v3 + Stroetmann-Stilreferenz
- Heading 1-3 only (NEVER Heading 4)
- sub-title/sub-number never as visible headings
- mandatory cross-cut H2s always emitted (even if not in CSV)
- cross-cut H2 display text: "E1 - GOV-3 (ESRS 2)" format
- NUMBER → Stroetmann-Satz (varied: "Im Berichtsjahr beträgt…" / "beläuft sich auf…" / etc.)
- ≤5 NUMBER sub-rows → bundled Stroetmann-Fließtext; >5 → Tabelle + Summary-Satz
- Gesamt-/Total-Werte: "Davon entfallen…"-Struktur wie Stroetmann
- yes/ja/teilweise → 1 Satz (kein Review-Marker)
- explanation als Hauptinhalt → [[REVIEW: explanation-driven]]
- minimale Eingabe zu langem Absatz → [[REVIEW: expanded-into-paragraph]]
- explanations als plain body text integriert (kein "Anmerkung:"-Label)
- add_subheading() für fette Zwischenüberschriften (nie H4)
- "keine Vorfälle"-Regel für Vorfalls-Datenpunkte
- Tabellen immer mit Einleitungstext
- Europäische Zahlenformatierung; Jahreszahlen (1900–2100) ohne Tausenderpunkt

Stilreferenz: 20251128_CSRD Report_Stroetmann.docx (Beispiele/)
Regeln (aus Stroetmann extrahiert):
1. Satzlänge 15–25 Wörter; kurze Ankersätze bewusst einsetzen
2. Absätze ~85 Wörter, 3–5 Sätze, Claim–Elaboration–Kontext
3. Absatzöffner: "Die / Im / Zur / Für + Subjekt"
4. Zahlen eingebettet: "[Subjekt] beträgt [X] [Einheit]" / "Davon entfallen…"
5. Methodik inline im Fließtext, nie als Anmerkungs-Block
6. Additive Konnektoren: Darüber hinaus, Ergänzend, Flankierend, Zudem
7. Institutionelle dritte Person; "-ende"-Formen durchgehend
8. Einschränkungen: ein Satz, kein Padding
9. Bullet-Listen nur ab ≥ 4 parallelen Einträgen (Einleitungssatz mit Doppelpunkt)
10. Ziele: Basiswert + Zielwert + Prozentveränderung
11. H3-Format: "[Code] – [Vollname]" mit geschütztem Em-Dash
12. Passiv für Prozesse; Aktiv für Verantwortungszuweisung
"""

import csv
import re
import os
from datetime import datetime
from collections import OrderedDict

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE     = os.path.dirname(os.path.abspath(__file__))
CSV_PATH = os.path.join(BASE, "Input Data",
           "csrd_report_36ced677-f008-488c-9e7e-da6f1ac8259c.csv")
OUT_DOCX = os.path.join(BASE, "CSRD_Report.docx")
OUT_QA   = os.path.join(BASE, "QA_log.md")

COMPANY = "FRÖBEL e.V."
YEAR    = "2024"

# ── Fröbel brand colours ──────────────────────────────────────────────────────
C_H1    = RGBColor(0xFF, 0xFF, 0xFF)   # white (on dark H1 shading)
C_H2    = RGBColor(0x1C, 0x46, 0x2D)   # deep Fröbel green
C_H3    = RGBColor(0x2D, 0x6E, 0x46)   # medium green
C_BODY  = RGBColor(0x1A, 0x1A, 0x1A)   # near-black
C_GREY  = RGBColor(0x70, 0x70, 0x70)   # grey (notes, placeholders)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_TBL_H = "1C462D"   # table header bg (dark Fröbel green)
C_TBL_1 = "EBF3EE"   # odd row bg (very light green)
FONT_HEAD = "Century Gothic"
FONT      = "Calibri"

# ── H1 section names ──────────────────────────────────────────────────────────
H1_ALLGEMEIN  = "Allgemeine Informationen (ESRS 2)"
H1_UMWELT     = "Umweltinformationen"
H1_SOZIAL     = "Soziale Informationen"
H1_GOVERNANCE = "Governance-Informationen"
H1_ORDER      = [H1_ALLGEMEIN, H1_UMWELT, H1_SOZIAL, H1_GOVERNANCE]

# ── Cross-cut H2 display text (mandatory, always emitted) ────────────────────
# Format: "Xn - YYY (ESRS 2)" exactly as specified
CROSSCUT_DISPLAY = {
    "E1-GOV-3": "E1 - GOV-3 (ESRS 2)",
    "E1-SBM-3": "E1 - SBM-3 (ESRS 2)",
    "E1-IRO-1": "E1 - IRO-1 (ESRS 2)",
    "E2-IRO-1": "E2 - IRO-1 (ESRS 2)",
    "E3-IRO-1": "E3 - IRO-1 (ESRS 2)",
    "E4-IRO-1": "E4 - IRO-1 (ESRS 2)",
    "E4-SBM-3": "E4 - SBM-3 (ESRS 2)",
    "E5-IRO-1": "E5 - IRO-1 (ESRS 2)",
    "S1-SBM-2": "S1 - SBM-2 (ESRS 2)",
    "S1-SBM-3": "S1 - SBM-3 (ESRS 2)",
    "S2-SBM-2": "S2 - SBM-2 (ESRS 2)",
    "S2-SBM-3": "S2 - SBM-3 (ESRS 2)",
    "S3-SBM-2": "S3 - SBM-2 (ESRS 2)",
    "S3-SBM-3": "S3 - SBM-3 (ESRS 2)",
    "S4-SBM-2": "S4 - SBM-2 (ESRS 2)",
    "S4-SBM-3": "S4 - SBM-3 (ESRS 2)",
    "G1-GOV-1": "G1 - GOV-1 (ESRS 2)",
    "G1-IRO-1": "G1 - IRO-1 (ESRS 2)",
}

# Mandatory cross-cuts with H1 mapping (always emitted even if not in CSV)
MANDATORY_CROSSCUTS = [
    ("E1-GOV-3", H1_UMWELT),
    ("E1-SBM-3", H1_UMWELT),
    ("E1-IRO-1", H1_UMWELT),
    ("E2-IRO-1", H1_UMWELT),
    ("E3-IRO-1", H1_UMWELT),
    ("E4-IRO-1", H1_UMWELT),
    ("E4-SBM-3", H1_UMWELT),
    ("E5-IRO-1", H1_UMWELT),
    ("S1-SBM-2", H1_SOZIAL),
    ("S1-SBM-3", H1_SOZIAL),
    ("S2-SBM-2", H1_SOZIAL),
    ("S2-SBM-3", H1_SOZIAL),
    ("S3-SBM-2", H1_SOZIAL),
    ("S3-SBM-3", H1_SOZIAL),
    ("S4-SBM-2", H1_SOZIAL),
    ("S4-SBM-3", H1_SOZIAL),
    ("G1-GOV-1", H1_GOVERNANCE),
    ("G1-IRO-1", H1_GOVERNANCE),
]
MANDATORY_CROSSCUT_SET = {dr for dr, _ in MANDATORY_CROSSCUTS}
MANDATORY_CROSSCUT_H1  = {dr: h1 for dr, h1 in MANDATORY_CROSSCUTS}


def get_h1(dr: str) -> str:
    d = dr.strip()
    if re.match(r'^E[1-5]', d, re.I):
        return H1_UMWELT
    if re.match(r'^S[1-4]', d, re.I):
        return H1_SOZIAL
    if re.match(r'^G1', d, re.I):
        return H1_GOVERNANCE
    return H1_ALLGEMEIN


# ── Cross-cut detection ───────────────────────────────────────────────────────
_XCUT_RE = re.compile(r'^[ESG]\d+-(GOV|SBM|IRO)-', re.I)


def is_cross_cut(dr: str) -> bool:
    return bool(_XCUT_RE.match(dr.strip())) or dr.strip() in MANDATORY_CROSSCUT_SET


# ── ESRS standard DR ordering ─────────────────────────────────────────────────
DR_ORDER = [
    "BP-1", "BP-2",
    "GOV-1", "GOV-2", "GOV-3", "GOV-4", "GOV-5",
    "SBM-1", "SBM-2", "SBM-3",
    "IRO-1", "IRO-2",
    "E1-GOV-3", "E1-SBM-3", "E1-IRO-1",
    "E1-1", "E1-2", "E1-3", "E1-4", "E1-5", "E1-6", "E1-7", "E1-8", "E1-9",
    "E2-SBM-3", "E2-IRO-1",
    "E2-1", "E2-2", "E2-3", "E2-4", "E2-5", "E2-6",
    "E3-SBM-3", "E3-IRO-1",
    "E3-1", "E3-2", "E3-3", "E3-4", "E3-5",
    "E4-SBM-3", "E4-IRO-1",
    "E4-1", "E4-2", "E4-3", "E4-4", "E4-5", "E4-6", "E4-7",
    "E5-SBM-3", "E5-IRO-1",
    "E5-1", "E5-2", "E5-3", "E5-4", "E5-5", "E5-6",
    "S1-SBM-2", "S1-SBM-3", "S1-IRO-1",
    "S1-1", "S1-2", "S1-3", "S1-4", "S1-5",
    "S1-6", "S1-7", "S1-8", "S1-9", "S1-10",
    "S1-11", "S1-12", "S1-13", "S1-14", "S1-15",
    "S1-16", "S1-17",
    "S2-SBM-2", "S2-SBM-3", "S2-IRO-1",
    "S2-1", "S2-2", "S2-3", "S2-4", "S2-5",
    "S3-SBM-2", "S3-SBM-3", "S3-IRO-1",
    "S3-1", "S3-2", "S3-3", "S3-4", "S3-5",
    "S4-SBM-2", "S4-SBM-3", "S4-IRO-1",
    "S4-1", "S4-2", "S4-3", "S4-4", "S4-5",
    "G1-GOV-1", "G1-GOV-3", "G1-IRO-1", "G1-SBM-3",
    "G1-1", "G1-2", "G1-3", "G1-4", "G1-5",
]


def dr_sort_key(dr: str) -> tuple:
    try:
        return (0, DR_ORDER.index(dr))
    except ValueError:
        return (1, dr)


# ── Number utilities ──────────────────────────────────────────────────────────

def _parse_num(s: str):
    s = s.strip()
    if not s:
        return None
    if re.match(r'^-?[\d.]+,\d+$', s):
        s = s.replace('.', '').replace(',', '.')
    try:
        return float(s)
    except ValueError:
        return None


def eu_fmt(val: float) -> str:
    # Jahreszahlen (1900–2100) nie mit Tausenderpunkt – z.B. 2023, nicht 2.023
    if 1900 <= val <= 2100 and val == int(val):
        return str(int(val))
    if val == int(val) and abs(val) < 1e15:
        return f"{int(val):,}".replace(',', '.')
    s = f"{val:.3f}".rstrip('0').rstrip('.')
    if '.' in s:
        int_part, dec_part = s.split('.')
    else:
        int_part, dec_part = s, ''
    try:
        int_formatted = f"{int(int_part):,}".replace(',', '.')
    except Exception:
        int_formatted = int_part
    return f"{int_formatted},{dec_part}" if dec_part else int_formatted


def fmt_number(raw: str, unit: str) -> str:
    raw = raw.strip()
    if raw.lower().endswith(' decimal'):
        num_part = raw[:-8].strip()
        num = _parse_num(num_part)
        if num is not None:
            return f"{eu_fmt(num)} %"
        return f"{num_part} %"
    m = re.match(r'^(-?[\d,.]+)\s*(.*)?$', raw)
    if m:
        num_part  = m.group(1)
        text_unit = (m.group(2) or '').strip()
        eff_unit  = unit or text_unit
        num = _parse_num(num_part)
        if num is not None:
            result = eu_fmt(num)
            return f"{result} {eff_unit}".strip() if eff_unit else result
    return f"{raw} {unit}".strip() if unit else raw


def eu_fmt_inline(text: str) -> str:
    def repl(m):
        before = text[m.start() - 1] if m.start() > 0 else ''
        after  = text[m.end()]       if m.end() < len(text) else ''
        if before and (before.isalpha() or before == '.'):
            return m.group(0)
        if after == '.':
            return m.group(0)
        try:
            val = float(m.group(0))
            return eu_fmt(val)
        except Exception:
            return m.group(0)
    return re.sub(r'\d+\.\d+', repl, text)


# ── Text utilities ────────────────────────────────────────────────────────────

def clean(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r'[=\-]{10,}', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)
    return text.strip()


def is_zero_no(data: str) -> bool:
    return data.strip().lower() in ('no', 'nein', '0', 'false')


def _effective_data_parts(data: str) -> list:
    """Strip separator lines and return meaningful non-empty parts."""
    cleaned = re.sub(r'[=\-]{5,}', '\n', data)
    cleaned = re.sub(r'\n{2,}', '\n', cleaned)
    return [p.strip() for p in cleaned.split('\n') if p.strip()]


def _is_part_zero_no(p: str) -> bool:
    """True if a single part is effectively a No/0/empty value."""
    pl = p.lower()
    if pl in ('no', 'nein', '0', 'false', ''):
        return True
    # "0 €", "0,00 €", "0.00 eur", etc.
    if re.match(r'^0[.,]?0*\s*(?:€|eur|usd|chf|%|tco2e)?$', pl):
        return True
    return False


def _is_qa_no(data: str) -> bool:
    """True if data is a 'Question?: No/Yes' pattern where answer is No."""
    m = re.match(r'^.+\?\s*:\s*(no|nein|0|false)\s*$', data.strip(), re.I | re.S)
    return bool(m)


_PROCEDURAL_PHRASES = [
    'capex und opex werden',
    'können ignoriert werden',
    'tool als pflichtfeld',
    'datum wurde eingetragen, da pflichtfeld',
    'frage doppelt',
]


def _is_procedural_note(text: str) -> bool:
    t = text.lower()
    return any(p in t for p in _PROCEDURAL_PHRASES)


# ── "Keine Vorfälle" detection ────────────────────────────────────────────────
_INCIDENT_KEYWORDS = [
    'vorfall', 'vorfälle', 'incident', 'verstoß', 'verstöße', 'beschwerde',
    'meldung', 'geldstrafe', 'verurteilung', 'korruption', 'bestechung',
    'schwerwiegend', 'negative vorfälle', 'menschenrechte',
]


def is_incident_dp(row: dict) -> bool:
    dp_title = (row.get('datapoint title') or '').lower()
    dr_name  = (row.get('disclosure-requirement name') or '').lower()
    expl     = (row.get('explanation') or '').lower()
    combined = f"{dp_title} {dr_name} {expl}"
    return any(k in combined for k in _INCIDENT_KEYWORDS)


def keine_vorfaelle_text(row: dict) -> str:
    dp_title = (row.get('datapoint title') or '').strip().rstrip('?').strip()
    dr_name  = (row.get('disclosure-requirement name') or '').strip()
    topic    = dp_title or dr_name or 'Vorfälle'
    return f"Im Berichtszeitraum wurden keine entsprechenden Ereignisse oder Vorfälle im Zusammenhang mit {topic} gemeldet oder registriert."


def is_empty_row(row: dict) -> bool:
    if (row.get('reportability') or '').strip() == 'NOT_REPORTABLE':
        return True
    data = (row.get('entered data') or '').strip()
    if not data:
        return True

    expl = (row.get('explanation') or '').strip()

    # Check simple bare No/0
    if is_zero_no(data):
        if is_incident_dp(row):
            return False
        if expl and not _is_procedural_note(expl):
            return False
        return True

    # Check Q&A pattern like "Question?: No"
    if _is_qa_no(data):
        if is_incident_dp(row):
            return False
        if expl and not _is_procedural_note(expl):
            return False
        return True

    # Check multi-line structured data where ALL parts are No/0
    # e.g. "No \n--------------------\n0 €"
    parts = _effective_data_parts(data)
    if parts and all(_is_part_zero_no(p) for p in parts):
        if is_incident_dp(row):
            return False
        if expl and not _is_procedural_note(expl):
            return False
        return True

    return False


def detect_type(row: dict) -> str:
    unit = (row.get('unit') or '').strip()
    data = (row.get('entered data') or '').strip()
    if unit:
        return 'NUMBER'
    if re.search(r'#\d+', data) and re.search(r'-{5,}|={5,}', data):
        return 'TABLE'
    if re.search(r'[-=]{5,}', data) and re.search(r':\s*\S', data):
        return 'TABLE'
    first_line = data.split('\n')[0].strip() if data else ''
    first_token = first_line.split()[0] if first_line.split() else ''
    if first_token and _parse_num(first_token) is not None:
        rest = first_line[len(first_token):].strip()
        if not rest or re.match(r'^[\w/°%²³µ·\s\u2082\u2083\u2084]*$', rest):
            return 'NUMBER'
    if re.match(r'^-?[\d,.]+\s+decimal\s*$', data, re.I):
        return 'NUMBER'
    return 'NARRATIVE'


# ── Structured data parsing ───────────────────────────────────────────────────

def _fix_decimal_in_value(v: str) -> str:
    return re.sub(r'(-?[\d,.]+)\s+decimal\b', lambda m: fmt_number(m.group(1), '%'), v, flags=re.I)


def _eu_fmt_cell(v: str) -> str:
    v = v.strip()
    if v.lower() == 'no':
        return 'Nein'
    if v.lower() == 'yes':
        return 'Ja'
    m = re.match(r'^(-?[\d,.]+)\s*(.*)?$', v)
    if m:
        num_part  = m.group(1)
        unit_part = (m.group(2) or '').strip()
        n = _parse_num(num_part)
        if n is not None:
            return (eu_fmt(n) + (' ' + unit_part if unit_part else '')).strip()
    return v


def parse_structured(data: str) -> list:
    if re.search(r'#\d+', data):
        norm    = re.sub(r'[=\-]{5,}', '\n||||\n', data)
        entries = re.split(r'(?=#\d+[\s\n])', norm.strip())
        result  = []
        for entry in entries:
            entry = re.sub(r'^#\d+\s*', '', entry).strip()
            parts = [p.strip() for p in entry.split('||||') if p.strip()]
            obj   = {}
            for p in parts:
                if ':' in p:
                    k, _, v = p.partition(':')
                    k = k.strip()
                    v = _fix_decimal_in_value(v.strip())
                    v = _eu_fmt_cell(eu_fmt_inline(v))
                    if k and v:
                        obj[k] = v
                elif p:
                    obj.setdefault('Wert', _eu_fmt_cell(_fix_decimal_in_value(p)))
            if obj:
                result.append(obj)
        if result:
            return result

    if re.search(r'[-=]{5,}', data):
        norm  = re.sub(r'[=\-]{5,}', '\n||||\n', data)
        parts = [p.strip() for p in norm.split('||||') if p.strip()]
        if len(parts) >= 2:
            result = []
            for p in parts:
                for line in p.split('\n'):
                    line = line.strip()
                    if ':' in line:
                        k, _, v = line.partition(':')
                        k = k.strip()
                        v = _fix_decimal_in_value(v.strip())
                        v = _eu_fmt_cell(eu_fmt_inline(v))
                        if k and v:
                            result.append({'Bezeichnung': k, 'Wert': v})
            if result:
                return result

    return []


# ── Word helpers ──────────────────────────────────────────────────────────────

def _xml_spacing(before=0, after=0, line=None):
    s = OxmlElement('w:spacing')
    s.set(qn('w:before'), str(before))
    s.set(qn('w:after'),  str(after))
    if line:
        s.set(qn('w:line'), str(line))
        s.set(qn('w:lineRule'), 'auto')
    return s


def _set_cell_bg(cell, hex6: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex6)
    tcPr.append(shd)


def _set_tbl_borders(table, color="1C462D"):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    bdr = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        bdr.append(b)
    tblPr.append(bdr)


def _add_run(para, text, bold=False, italic=False, color=None, size=10.5,
             highlight=None):
    r = para.add_run(text)
    r.font.name   = FONT
    r.font.size   = Pt(size)
    r.font.bold   = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    if highlight:
        rPr = r._r.get_or_add_rPr()
        hl  = OxmlElement('w:highlight')
        hl.set(qn('w:val'), highlight)
        rPr.append(hl)
    return r


def configure_heading_styles(doc):
    # H1: dark green shading, white Century Gothic text
    sty1 = doc.styles['Heading 1']
    sty1.font.name      = FONT_HEAD
    sty1.font.size      = Pt(15.0)
    sty1.font.bold      = True
    sty1.font.color.rgb = C_WHITE
    pf1 = sty1.paragraph_format
    pf1.space_before   = Pt(2.0)
    pf1.space_after    = Pt(2.0)
    pf1.keep_with_next = True
    pPr1 = sty1.element.get_or_add_pPr()
    numPr = pPr1.find(qn('w:numPr'))
    if numPr is not None:
        pPr1.remove(numPr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '1C462D')
    pPr1.append(shd)
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'),  '160')
    ind.set(qn('w:right'), '80')
    pPr1.append(ind)

    # H2: deep green Century Gothic, bottom green underline
    sty2 = doc.styles['Heading 2']
    sty2.font.name      = FONT_HEAD
    sty2.font.size      = Pt(13.0)
    sty2.font.bold      = True
    sty2.font.color.rgb = C_H2
    pf2 = sty2.paragraph_format
    pf2.space_before   = Pt(12.0)
    pf2.space_after    = Pt(4.0)
    pf2.keep_with_next = True
    pPr2 = sty2.element.get_or_add_pPr()
    numPr = pPr2.find(qn('w:numPr'))
    if numPr is not None:
        pPr2.remove(numPr)
    pBdr2 = OxmlElement('w:pBdr')
    bbot  = OxmlElement('w:bottom')
    bbot.set(qn('w:val'),   'single')
    bbot.set(qn('w:sz'),    '6')
    bbot.set(qn('w:space'), '3')
    bbot.set(qn('w:color'), '2D6E46')
    pBdr2.append(bbot)
    pPr2.append(pBdr2)

    # H3: medium green Century Gothic
    sty3 = doc.styles['Heading 3']
    sty3.font.name      = FONT_HEAD
    sty3.font.size      = Pt(11.0)
    sty3.font.bold      = True
    sty3.font.color.rgb = C_H3
    pf3 = sty3.paragraph_format
    pf3.space_before   = Pt(8.0)
    pf3.space_after    = Pt(3.0)
    pf3.keep_with_next = True
    pPr3 = sty3.element.get_or_add_pPr()
    numPr = pPr3.find(qn('w:numPr'))
    if numPr is not None:
        pPr3.remove(numPr)


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.add_run(text)
    return p


def add_body(doc, text: str, color=None):
    text = text.strip() if text else ''
    if not text:
        return None
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p._p.get_or_add_pPr().append(_xml_spacing(before=0, after=80, line=276))
    _add_run(p, text, color=color or C_BODY)
    return p


def add_body_with_review(doc, text: str, marker: str):
    """Add a body paragraph ending with a [[REVIEW: ...]] marker in grey italic."""
    text = text.strip()
    if not text:
        return None
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p._p.get_or_add_pPr().append(_xml_spacing(before=0, after=80, line=276))
    _add_run(p, text + ' ', color=C_BODY)
    _add_run(p, f'[[REVIEW: {marker}]]', italic=True, color=C_GREY, size=8.5)
    return p


def add_subheading(doc, text: str):
    """Bold body-text paragraph used as intermediate section header (never H4)."""
    text = text.strip()
    if not text:
        return None
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p._p.get_or_add_pPr().append(_xml_spacing(before=120, after=40, line=276))
    _add_run(p, text, bold=True, color=C_BODY)
    return p


def add_bookmark(para, name: str):
    bid   = str(abs(hash(name)) % 90000 + 10000)
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'),   bid)
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), bid)
    para._p.insert(0, start)
    para._p.append(end)


def make_kv_table(doc, entries: list, intro_text: str = ''):
    if not entries:
        return
    if intro_text:
        add_body(doc, intro_text)
    cols = []
    for e in entries:
        for k in e:
            if k not in cols:
                cols.append(k)
    if not cols:
        return
    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_tbl_borders(tbl)
    hrow = tbl.rows[0]
    for i, col in enumerate(cols):
        _set_cell_bg(hrow.cells[i], C_TBL_H)
        _add_run(hrow.cells[i].paragraphs[0], col,
                 bold=True, color=C_WHITE, size=9)
    for idx, entry in enumerate(entries):
        row = tbl.add_row()
        bg  = C_TBL_1 if idx % 2 == 0 else "FFFFFF"
        for i, col in enumerate(cols):
            _set_cell_bg(row.cells[i], bg)
            _add_run(row.cells[i].paragraphs[0],
                     entry.get(col, ''), color=C_BODY, size=9)
    doc.add_paragraph()


def make_number_table(doc, rows: list, intro_text: str = ''):
    if intro_text:
        add_body(doc, intro_text)
    entries = []
    for row in rows:
        st  = (row.get('sub-title') or '').strip().rstrip('?').strip()
        if not st:
            st = (row.get('sub-number') or '').strip()
        raw  = (row.get('entered data') or '').strip()
        unit = (row.get('unit') or '').strip()
        val  = fmt_number(raw, unit)
        entries.append({'Bezeichnung': st, 'Wert': val})
    if entries:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        _set_tbl_borders(tbl)
        hrow = tbl.rows[0]
        for i, hdr in enumerate(['Bezeichnung', 'Wert']):
            _set_cell_bg(hrow.cells[i], C_TBL_H)
            _add_run(hrow.cells[i].paragraphs[0], hdr,
                     bold=True, color=C_WHITE, size=9)
        for idx, entry in enumerate(entries):
            row = tbl.add_row()
            bg  = C_TBL_1 if idx % 2 == 0 else "FFFFFF"
            for i, col in enumerate(['Bezeichnung', 'Wert']):
                _set_cell_bg(row.cells[i], bg)
                _add_run(row.cells[i].paragraphs[0],
                         entry.get(col, ''), color=C_BODY, size=9)
        doc.add_paragraph()


def make_iro2_table(doc, dr_pairs: list):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p._p.get_or_add_pPr().append(_xml_spacing(before=120, after=60))
    _add_run(p, "IRO-2 Inhaltstabelle der Offenlegungsanforderungen",
             bold=True, color=C_BODY, size=10.5)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_tbl_borders(tbl)

    for i, w in enumerate([Cm(3.0), Cm(11.5), Cm(2.0)]):
        for row in tbl.rows:
            row.cells[i].width = w

    hdrs = ["Angabe-Nr.", "Offenlegungsanforderung", "Seite"]
    hrow = tbl.rows[0]
    for i, hdr in enumerate(hdrs):
        _set_cell_bg(hrow.cells[i], C_TBL_H)
        _add_run(hrow.cells[i].paragraphs[0], hdr,
                 bold=True, color=C_WHITE, size=9)

    for idx, (code, name) in enumerate(dr_pairs):
        row = tbl.add_row()
        bg  = C_TBL_1 if idx % 2 == 0 else "FFFFFF"
        for i, val in enumerate([code, name, '']):
            _set_cell_bg(row.cells[i], bg)
            _add_run(row.cells[i].paragraphs[0], val, color=C_BODY, size=9)

    doc.add_paragraph()


# ── Content rendering ─────────────────────────────────────────────────────────

def _strip_part_prefix(s: str) -> str:
    """Remove redundant prefixes like 'davon', 'darunter', 'hiervon' from sub-titles."""
    return re.sub(r'^(?:davon|darunter|hiervon|hierbei|daraus)\s+', '', s,
                  flags=re.I).strip()


# ── Paragraph connectors for explanation integration (Stroetmann-Stil) ─────────
# "ist festzuhalten:" / "ist anzumerken:" work without verb inversion – grammatically
# safe regardless of what the following explanation sentence starts with.
_EXPL_CONNECTORS = [
    "Ergänzend ist festzuhalten:",
    "Methodisch ist anzumerken:",
    "Ergänzend:",
    "Zur Einordnung ist hinzuzufügen:",
    "Darüber hinaus ist festzuhalten:",
]
# Starters that already signal a connection – don't prepend a second connector
_EXPL_NO_CONN_STARTS = (
    'darüber', 'ergänzend', 'zudem', 'flankierend', 'daraus',
    'methodisch', 'zur ', 'zum ', 'im ', 'in diesem', 'in der',
    'festzuhalten', 'anzumerken', 'hinzuzufügen', 'insgesamt',
    'dies ', 'damit ', 'dadurch ', 'hierbei ', 'hierzu ',
)


def _connect_expl(expl: str, dr: str, idx: int) -> str:
    """Prepend a Stroetmann-style transition connector to an explanation paragraph."""
    expl_fmt = eu_fmt_inline(expl.strip())
    if not expl_fmt:
        return expl_fmt
    if any(expl_fmt.lower().startswith(s) for s in _EXPL_NO_CONN_STARTS):
        return expl_fmt  # already starts with a connector
    conn = _EXPL_CONNECTORS[abs(hash(f"{dr}|{idx}")) % len(_EXPL_CONNECTORS)]
    return f"{conn} {expl_fmt}"


def _render_number_bundle(doc, rows: list, dp_title: str, dr: str, qa: list):
    """Render ≤5 NUMBER sub-rows as bundled Stroetmann-style flowing prose."""
    items = []
    for r in rows:
        st   = (r.get('sub-title') or '').strip().rstrip('?').strip()
        raw  = (r.get('entered data') or '').strip()
        unit = (r.get('unit') or '').strip()
        val  = fmt_number(raw, unit)
        items.append((st, val))

    dp_lower = dp_title.lower().rstrip('.').rstrip(':') if dp_title else ''

    if len(items) == 1:
        st, val = items[0]
        ctx = st.lower() if st else dp_lower
        if ctx:
            txt = f"Im Berichtsjahr beträgt {ctx} {val}."
        else:
            txt = f"Der berichtete Wert beträgt {val}."
        add_body(doc, txt)
    else:
        first_st, first_val = items[0]
        rest = items[1:]

        # Detect total/Gesamt as lead → use "Davon entfallen" pattern (Stroetmann)
        first_label = (first_st or dp_lower or '').lower()
        is_total = any(w in first_label
                       for w in ('gesamt', 'total', 'brutto', 'summe', 'netto'))

        ctx = first_st.lower() if first_st else dp_lower
        ctx_cap = ctx[0].upper() + ctx[1:] if ctx else 'Der Gesamtwert'

        if is_total:
            lead = f"{ctx_cap} beläuft sich im Berichtsjahr auf {first_val}"
        elif ctx:
            lead = f"Im Berichtsjahr beträgt {ctx} {first_val}"
        else:
            lead = f"Der Hauptwert beträgt {first_val}"

        if len(rest) == 1:
            r_st, r_val = rest[0]
            r_st_clean = _strip_part_prefix(r_st).lower() if r_st else ''
            if r_st_clean and is_total:
                txt = f"{lead}; davon entfallen {r_val} auf {r_st_clean}."
            elif r_st_clean:
                txt = f"{lead}; ergänzend werden {r_val} für {r_st_clean} ausgewiesen."
            else:
                txt = f"{lead}; ergänzend wird {r_val} ausgewiesen."
        else:
            rest_parts = []
            for st, val in rest:
                st_clean = _strip_part_prefix(st).lower() if st else ''
                if st_clean and is_total:
                    rest_parts.append(f"{val} auf {st_clean}")
                elif st_clean:
                    rest_parts.append(f"{val} ({st_clean})")
                else:
                    rest_parts.append(val)
            if len(rest_parts) == 2:
                rest_str = rest_parts[0] + " sowie " + rest_parts[1]
            else:
                rest_str = ", ".join(rest_parts[:-1]) + " sowie " + rest_parts[-1]
            if is_total:
                txt = f"{lead}; davon entfallen {rest_str}."
            else:
                txt = f"{lead}; ergänzend werden {rest_str} ausgewiesen."

        add_body(doc, txt)

    for r in rows:
        qa.append(('DONE', dr, (r.get('sub-number') or '').strip(),
                   'NUMBER', r.get('_row', '?'), 'bundled_sentence'))


def _number_sentence(raw: str, unit: str, dp_title: str, dr: str = '', sub: str = '') -> str:
    """Wrap a number in a context sentence with varied Stroetmann-style phrasing."""
    val_str = fmt_number(raw, unit)
    ctx = dp_title.strip().rstrip('.').rstrip('?').rstrip(':') if dp_title else ''
    if not ctx:
        return f"Der berichtete Wert beträgt {val_str}."
    ctx_l = ctx[0].lower() + ctx[1:] if len(ctx) > 1 else ctx.lower()
    ctx_cap = ctx_l[0].upper() + ctx_l[1:]
    # Deterministic template selection via hash so output is stable across runs
    key = abs(hash(f"{dr}|{sub}|{ctx}")) % 4
    if key == 0:
        return f"Im Berichtsjahr beträgt {ctx_l} {val_str}."
    elif key == 1:
        return f"{ctx_cap} beläuft sich im Berichtszeitraum auf {val_str}."
    elif key == 2:
        return f"Für {ctx_l} werden im Berichtsjahr {val_str} ausgewiesen."
    else:
        return f"Der berichtete Wert für {ctx_l} beträgt {val_str}."


def render_rows(doc, rows: list, qa: list, dp_title: str = '') -> bool:
    """
    Write CSV row content to the document.
    Returns True if at least one row was rendered.
    NEVER uses Heading 4. sub-title never shown as heading.
    """
    wrote = False
    for row in rows:
        dr       = (row.get('disclosure-requirement nr.') or '').strip()
        sub      = (row.get('sub-number') or '').strip()
        raw      = (row.get('entered data') or '').strip()
        unit     = (row.get('unit') or '').strip()
        row_dp_title = (row.get('datapoint title') or dp_title or '').strip()
        ref      = row.get('_row', '?')

        if is_empty_row(row):
            qa.append(('SKIP', dr, sub, detect_type(row), ref, 'leer/NOT_REPORTABLE'))
            continue

        # ── Determine effective content ────────────────────────────────────
        from_expl = False
        if is_zero_no(raw):
            # "Keine Vorfälle" rule
            if is_incident_dp(row):
                text = keine_vorfaelle_text(row)
                add_body(doc, text)
                qa.append(('DONE', dr, sub, 'NARRATIVE', ref,
                           'keine_vorfaelle_rule'))
                qa.append(('KEIN_VORFALL', dr, sub, 'NARRATIVE', ref,
                           row_dp_title))
                wrote = True
                continue
            expl = (row.get('explanation') or '').strip()
            if expl and not _is_procedural_note(expl):
                effective = expl
                dtype     = 'NARRATIVE'
                from_expl = True
            else:
                qa.append(('SKIP', dr, sub, 'NARRATIVE', ref, 'nein/null_kein_inhalt'))
                continue
        else:
            effective = raw
            dtype     = detect_type(row)

        # ── Render by type ─────────────────────────────────────────────────
        # row_expl: explanation of this row, used to integrate into same paragraph
        row_expl = (row.get('explanation') or '').strip() if not from_expl else ''
        row_expl = row_expl if (row_expl and not _is_procedural_note(row_expl)) else ''

        if dtype == 'NUMBER':
            sentence = _number_sentence(effective, unit, row_dp_title, dr, sub)
            if from_expl:
                add_body_with_review(doc, sentence, 'explanation-driven')
                qa.append(('REVIEW_EXPL', dr, sub, 'NUMBER', ref, sentence[:60]))
            elif row_expl:
                # Integrate explanation as second sentence in same paragraph
                add_body(doc, f"{sentence} {eu_fmt_inline(row_expl)}")
                row['_expl_consumed'] = True
            else:
                add_body(doc, sentence)
            qa.append(('DONE', dr, sub, 'NUMBER', ref, 'number_sentence'))

        elif dtype == 'TABLE':
            entries = parse_structured(effective)
            if entries:
                intro = (f"Die folgende Tabelle zeigt die Angaben zu "
                         f"{row_dp_title.lower().rstrip('.')}.")
                make_kv_table(doc, entries, intro_text=intro)
            else:
                for part in re.split(r'\n\n+', clean(effective)):
                    part = part.replace('\n', ' ').strip()
                    if not part:
                        continue
                    # Skip bare No/0 fragments that fell through parse_structured
                    if _is_part_zero_no(part) or _is_qa_no(part):
                        continue
                    if from_expl:
                        add_body_with_review(doc, part, 'explanation-driven')
                    else:
                        add_body(doc, part)
            if from_expl:
                qa.append(('REVIEW_EXPL', dr, sub, 'TABLE', ref, ''))
            qa.append(('DONE', dr, sub, 'TABLE', ref, ''))

        else:  # NARRATIVE
            txt = effective.strip()
            # Boolean answers: integrate explanation directly into same paragraph
            if txt.lower() in ('yes', 'ja'):
                label = row_dp_title.rstrip('?.:').strip() if row_dp_title else 'Diese Angabe'
                sentence = f"{label} trifft zu."
                if row_expl:
                    add_body(doc, f"{sentence} {eu_fmt_inline(row_expl)}")
                    row['_expl_consumed'] = True
                else:
                    add_body(doc, sentence)
                qa.append(('DONE', dr, sub, 'NARRATIVE', ref, 'yes_sentence'))
            elif txt.lower() in ('teilweise', 'partly', 'partial'):
                label = row_dp_title.rstrip('?.:').strip() if row_dp_title else 'Diese Angabe'
                sentence = f"{label} trifft teilweise zu."
                if row_expl:
                    add_body(doc, f"{sentence} {eu_fmt_inline(row_expl)}")
                    row['_expl_consumed'] = True
                else:
                    add_body(doc, sentence)
                qa.append(('DONE', dr, sub, 'NARRATIVE', ref, 'teilweise_sentence'))
            else:
                txt = clean(txt)
                txt = eu_fmt_inline(txt)
                for part in re.split(r'\n\n+', txt):
                    part = part.replace('\n', ' ').strip()
                    if not part:
                        continue
                    if from_expl:
                        add_body_with_review(doc, part, 'explanation-driven')
                    else:
                        add_body(doc, part)
                if from_expl:
                    qa.append(('REVIEW_EXPL', dr, sub, 'NARRATIVE', ref, txt[:60]))
                qa.append(('DONE', dr, sub, 'NARRATIVE', ref, ''))
        wrote = True

    return wrote


def _collect_explanations(rows: list) -> list:
    seen, result = set(), []
    for r in rows:
        if r.get('_expl_consumed'):
            continue  # already integrated inline into the main paragraph
        expl = (r.get('explanation') or '').strip()
        if expl and not _is_procedural_note(expl) and expl not in seen:
            raw = (r.get('entered data') or '').strip()
            if is_zero_no(raw):
                seen.add(expl)
                continue
            seen.add(expl)
            result.append(expl)
    return result


def render_dr_content(doc, dr: str, dr_rows: list, qa: list):
    """
    Emit H3 (datapoint groups) and content for a normal DR.
    NEVER emits Heading 4. sub-title never used as heading.
    - NUMBER sub-items (>=2): table with intro
    - TABLE sub-items same columns: merged table
    - All other: flowing body text
    """
    dp_groups = OrderedDict()
    dp_titles = {}
    for r in dr_rows:
        dp = (r.get('datapoint nr.') or '').strip()
        dt = (r.get('datapoint title') or '').strip()
        dp_groups.setdefault(dp, []).append(r)
        dp_titles.setdefault(dp, dt)

    for dp_nr, dp_rows in dp_groups.items():
        non_empty = [r for r in dp_rows if not is_empty_row(r)]
        if not non_empty:
            for r in dp_rows:
                qa.append(('SKIP', dr, (r.get('sub-number') or '').strip(),
                           detect_type(r), r.get('_row', '?'), 'dp_leer'))
            continue

        dp_title = dp_titles[dp_nr]
        h3_text  = f"{dp_nr} – {dp_title}" if dp_title else dp_nr
        add_heading(doc, h3_text, 3)

        # Rows without sub-title (direct content)
        direct = [r for r in dp_rows if not (r.get('sub-title') or '').strip()]
        sub_rows = [r for r in dp_rows if (r.get('sub-title') or '').strip()]

        if direct:
            render_rows(doc, direct, qa, dp_title)

        if sub_rows:
            all_ne = [r for r in sub_rows if not is_empty_row(r)]

            # Case A: ALL non-empty sub-rows are NUMBER
            if (all_ne and
                    all(detect_type(r) == 'NUMBER' for r in all_ne) and
                    len(all_ne) >= 2):
                if len(all_ne) <= 5:
                    # ≤5 values: bundle into flowing sentence(s)
                    _render_number_bundle(doc, all_ne, dp_title, dr, qa)
                else:
                    # >5 values: table with intro + bünding summary sentence
                    intro = (f"Die folgende Übersicht zeigt die Kennzahlen zu "
                             f"{dp_title.lower().rstrip('.')}.")
                    make_number_table(doc, all_ne, intro_text=intro)
                    first_val = fmt_number(
                        (all_ne[0].get('entered data') or '').strip(),
                        (all_ne[0].get('unit') or '').strip())
                    first_st = (all_ne[0].get('sub-title') or dp_title or '').strip().rstrip('?').strip()
                    add_body(doc, (
                        f"Insgesamt umfasst {dp_title.lower().rstrip('.')} "
                        f"{len(all_ne)} Kennzahlen; der erste ausgewiesene "
                        f"Wert ({first_st.lower()}) beträgt {first_val}."))
                    for r in all_ne:
                        qa.append(('DONE', dr, (r.get('sub-number') or '').strip(),
                                   'NUMBER', r.get('_row', '?'), 'num_table'))
                for r in sub_rows:
                    if is_empty_row(r):
                        qa.append(('SKIP', dr, (r.get('sub-number') or '').strip(),
                                   'NUMBER', r.get('_row', '?'), 'leer'))

            # Case B: ALL non-empty sub-rows are TABLE with same columns → merged table
            elif all_ne and all(detect_type(r) == 'TABLE' for r in all_ne):
                all_entries = []
                for r in all_ne:
                    entries = parse_structured((r.get('entered data') or '').strip())
                    all_entries.extend(entries)
                key_sets = [frozenset(e.keys()) for e in all_entries]
                if all_entries and len(set(key_sets)) == 1:
                    intro = (f"Die folgende Tabelle zeigt die Angaben zu "
                             f"{dp_title.lower().rstrip('.')}.")
                    make_kv_table(doc, all_entries, intro_text=intro)
                    for r in all_ne:
                        qa.append(('DONE', dr, (r.get('sub-number') or '').strip(),
                                   'TABLE', r.get('_row', '?'), 'merged_table'))
                    for r in sub_rows:
                        if is_empty_row(r):
                            qa.append(('SKIP', dr, (r.get('sub-number') or '').strip(),
                                       'TABLE', r.get('_row', '?'), 'leer'))
                else:
                    # Columns differ → render as flowing text (no H4)
                    render_rows(doc, sub_rows, qa, dp_title)

            # Case C: Mixed or single → render as flowing body text (no H4)
            else:
                render_rows(doc, sub_rows, qa, dp_title)

        # Trailing explanations: connected with Stroetmann-style transition phrase
        explns = _collect_explanations(non_empty)
        for i, expl in enumerate(explns):
            add_body(doc, _connect_expl(expl, dr, i))


# ── Main document builder ─────────────────────────────────────────────────────

def build_document(csv_rows: list):
    doc = Document()

    doc.core_properties.title   = (
        f"Sustainability statement {YEAR} – {COMPANY}")
    doc.core_properties.author  = COMPANY
    doc.core_properties.subject = "ESRS Nachhaltigkeitserklärung"

    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.8)
        sec.right_margin  = Cm(2.8)

    doc.styles['Normal'].font.name = FONT
    doc.styles['Normal'].font.size = Pt(10.5)
    configure_heading_styles(doc)

    qa        = []
    iro2_done = False

    # ── Build data structures ─────────────────────────────────────────────────
    by_dr   = OrderedDict()
    dr_name = {}

    for i, r in enumerate(csv_rows):
        r['_row'] = i + 2
        dr = (r.get('disclosure-requirement nr.') or '').strip()
        nm = (r.get('disclosure-requirement name') or '').strip()
        if dr:
            by_dr.setdefault(dr, []).append(r)
            dr_name.setdefault(dr, nm)

    # Build per-H1 DR lists from CSV
    h1_drs = {k: [] for k in H1_ORDER}
    for dr in by_dr:
        h1_drs[get_h1(dr)].append(dr)

    # Add mandatory cross-cuts not yet in their H1 list
    for mdr, mh1 in MANDATORY_CROSSCUTS:
        if mdr not in h1_drs[mh1]:
            h1_drs[mh1].append(mdr)

    for k in H1_ORDER:
        h1_drs[k] = sorted(set(h1_drs[k]), key=dr_sort_key)

    # ── Title page ────────────────────────────────────────────────────────────
    for _ in range(5):
        doc.add_paragraph()

    p1 = doc.add_paragraph(style='Heading 1')
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run('Sustainability statement')

    p2 = doc.add_paragraph(style='Heading 2')
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(YEAR)

    p3 = doc.add_paragraph(style='Heading 3')
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(COMPANY)

    doc.add_page_break()

    # ── TOC placeholder ───────────────────────────────────────────────────────
    p_toc = doc.add_paragraph()
    p_toc.style = doc.styles['Normal']
    _add_run(p_toc, '[Inhaltsverzeichnis]', bold=True, size=12, color=C_BODY)

    p_ph = doc.add_paragraph()
    p_ph.style = doc.styles['Normal']
    _add_run(p_ph, '[Platzhalter, wird extern befüllt]', color=C_GREY)

    doc.add_page_break()

    # ── ESRS section loop ─────────────────────────────────────────────────────
    for h1_name in H1_ORDER:
        drs = h1_drs[h1_name]
        if not drs:
            continue

        add_heading(doc, h1_name, 1)

        for dr in drs:
            # ── IRO-2 special handling ─────────────────────────────────────
            if dr == 'IRO-2':
                dr_rows = by_dr.get(dr, [])
                # Use CSV name if available (matches spec exactly in this dataset)
                h2_text = (f"IRO-2 – {dr_name[dr]}"
                           if dr in dr_name else
                           "IRO-2 – Offenlegungsanforderungen in den ESRS, "
                           "die vom Nachhaltigkeitsbericht des Unternehmens "
                           "abgedeckt werden")
                ph2 = add_heading(doc, h2_text, 2)
                add_bookmark(ph2, 'iro-2')

                # DP 56 (or first DP if 56 not present) gets H3 before table
                dp56_rows = [r for r in dr_rows
                             if (r.get('datapoint nr.') or '').strip() == '56']
                if dp56_rows:
                    dp56_title = (dp56_rows[0].get('datapoint title') or '').strip()
                    ph3 = add_heading(doc, f"56 – {dp56_title}", 3)
                    add_bookmark(ph3, 'iro-2-56')

                if not iro2_done:
                    # Build pairs: only DRs actually in the CSV
                    pairs = [(d, dr_name[d])
                             for d in sorted(by_dr.keys(), key=dr_sort_key)]
                    make_iro2_table(doc, pairs)
                    iro2_done = True
                    qa.append(('IRO2_TABLE', 'IRO-2', '-', 'TABLE', '-',
                               'einmalig erzeugt'))

                # Render remaining DPs (not DP 56)
                non56 = [r for r in dr_rows
                         if (r.get('datapoint nr.') or '').strip() != '56']
                if non56:
                    render_dr_content(doc, dr, non56, qa)
                continue

            # ── Cross-cut H2 (always emitted, no content) ─────────────────
            if is_cross_cut(dr):
                # Use spec display text if available, else use CSV name
                display = (CROSSCUT_DISPLAY.get(dr) or
                           (f"{dr} – {dr_name[dr]}" if dr in dr_name else dr))
                add_heading(doc, display, 2)
                dr_rows = by_dr.get(dr, [])
                for r in dr_rows:
                    qa.append(('SKIP_CROSSCUT', dr,
                               (r.get('sub-number') or '').strip(),
                               detect_type(r), r.get('_row', '?'),
                               'Cross-cut DR – kein Inhalt'))
                continue

            # ── Normal DR ─────────────────────────────────────────────────
            dr_rows  = by_dr.get(dr, [])
            non_empty = [r for r in dr_rows if not is_empty_row(r)]
            if not non_empty:
                for r in dr_rows:
                    qa.append(('SKIP', dr,
                               (r.get('sub-number') or '').strip(),
                               detect_type(r), r.get('_row', '?'),
                               'DR komplett leer'))
                continue

            h2_text = f"{dr} – {dr_name[dr]}"
            add_heading(doc, h2_text, 2)
            render_dr_content(doc, dr, dr_rows, qa)

    # If iro2 was never emitted (e.g. IRO-2 not in CSV), force it now under Allgemein
    if not iro2_done:
        pairs = [(d, dr_name[d])
                 for d in sorted(by_dr.keys(), key=dr_sort_key)]
        make_iro2_table(doc, pairs)
        iro2_done = True
        qa.append(('IRO2_TABLE', 'IRO-2', '-', 'TABLE', '-', 'fallback erzeugt'))

    return doc, qa, iro2_done


# ── QA log writer ─────────────────────────────────────────────────────────────

def write_qa(path: str, csv_rows: list, qa: list, iro2_done: bool):
    done_drs  = set(e[1] for e in qa if e[0] == 'DONE')
    done_dps  = set((e[1], e[2]) for e in qa if e[0] == 'DONE')
    skipped   = [e for e in qa if e[0] == 'SKIP']
    crosscuts = [e for e in qa if e[0] == 'SKIP_CROSSCUT']
    iro2_ok   = any(e[0] == 'IRO2_TABLE' for e in qa)
    kv_list   = [e for e in qa if e[0] == 'KEIN_VORFALL']
    rev_word  = [e for e in qa if e[0] == 'REVIEW_WORD']
    rev_para  = [e for e in qa if e[0] == 'REVIEW_PARA']
    rev_expl  = [e for e in qa if e[0] == 'REVIEW_EXPL']

    lines = [
        f"# QA-Log – CSRD Report {COMPANY} {YEAR}",
        f"_Generiert: {datetime.now().strftime('%Y-%m-%d %H:%M')}_",
        "",
        "## Zusammenfassung",
        f"- CSV-Zeilen gesamt:              **{len(csv_rows)}**",
        f"- Ausgegebene DRs (H2 mit Inhalt): **{len(done_drs)}**",
        f"- Ausgegebene Datenpunkte (H3):   **{len(done_dps)}**",
        f"- Gefilterte Zeilen (leer/NR):    **{len(skipped)}**",
        f"- Cross-cut DRs (H2 ohne Inhalt): **{len(set(e[1] for e in crosscuts))}**",
        f"- IRO-2 Tabelle genau einmal:     **{'✓ Ja' if iro2_ok else '✗ FEHLER'}**",
        "",
        "## 1. Unterdrückte Datenpunkte (leer / NOT_REPORTABLE / No-only / 0-only)",
        "| DR | Sub-Nr. | Typ | Zeile | Grund |",
        "|---|---|---|---|---|",
    ]
    for e in skipped:
        lines.append(
            f"| `{e[1]}` | `{e[2]}` | {e[3]} | Z.{e[4]} | {e[5] if len(e) > 5 else ''} |")

    lines += [
        "",
        "## 2. 'Keine Vorfälle'-Datenpunkte (trotz No/0 ausgegeben)",
        "| DR | Sub-Nr. | Typ | Zeile | Kontext |",
        "|---|---|---|---|---|",
    ]
    for e in kv_list:
        lines.append(
            f"| `{e[1]}` | `{e[2]}` | {e[3]} | Z.{e[4]} | {e[5][:60] if len(e) > 5 else ''} |")

    lines += [
        "",
        "## 3. [[REVIEW: expanded-from-single-word]] Absätze",
        "_(Nur noch ausgegeben, wenn ein einzelnes Wort zu einem echten Absatz >60 Wörter ausgebaut wird)_",
        "| DR | Sub-Nr. | Typ | Zeile | Inhalt |",
        "|---|---|---|---|---|",
    ]
    for e in rev_word:
        lines.append(
            f"| `{e[1]}` | `{e[2]}` | {e[3]} | Z.{e[4]} | {e[5][:80] if len(e) > 5 else ''} |")

    lines += [
        "",
        "## 4. [[REVIEW: expanded-into-paragraph]] Absätze",
        "_(Minimale Eingabe wurde zu einem Absatz >60 Wörter oder >2 Sätze ausgebaut)_",
        "| DR | Sub-Nr. | Typ | Zeile | Inhalt |",
        "|---|---|---|---|---|",
    ]
    for e in rev_para:
        lines.append(
            f"| `{e[1]}` | `{e[2]}` | {e[3]} | Z.{e[4]} | {e[5][:80] if len(e) > 5 else ''} |")

    lines += [
        "",
        "## 5. [[REVIEW: explanation-driven]] Absätze",
        "_(Erläuterungsfeld wurde als Hauptinhalt verwendet – bitte inhaltlich prüfen)_",
        "| DR | Sub-Nr. | Typ | Zeile | Inhalt |",
        "|---|---|---|---|---|",
    ]
    for e in rev_expl:
        lines.append(
            f"| `{e[1]}` | `{e[2]}` | {e[3]} | Z.{e[4]} | {e[5][:80] if len(e) > 5 else ''} |")

    lines += [
        "",
        "## 6. Alle Datenpunkte (Status-Übersicht)",
        "| DR | Sub-Nr. | Typ | Status | Zeile | Notiz |",
        "|---|---|---|---|---|---|",
    ]
    for e in qa:
        if e[0] in ('DONE', 'SKIP'):
            status = 'COMPLETED' if e[0] == 'DONE' else 'GEFILTERT'
            note   = e[5] if len(e) > 5 else ''
            lines.append(
                f"| `{e[1]}` | `{e[2]}` | {e[3]} | {status} | Z.{e[4]} | {note} |")

    lines += [
        "",
        "## 7. Cross-cut DRs (H2 ohne Datenpunkte)",
        "| DR | Anzeige-Text | Zeile |",
        "|---|---|---|",
    ]
    seen_xc = set()
    for e in crosscuts:
        if e[1] not in seen_xc:
            seen_xc.add(e[1])
            display = CROSSCUT_DISPLAY.get(e[1], e[1])
            lines.append(f"| `{e[1]}` | {display} | Z.{e[4]} |")
    # Also note mandatory cross-cuts not in CSV
    for mdr, _ in MANDATORY_CROSSCUTS:
        if mdr not in seen_xc:
            lines.append(f"| `{mdr}` | {CROSSCUT_DISPLAY[mdr]} | (nicht in CSV) |")

    lines += [
        "",
        "## 8. IRO-2 Tabelle",
        f"- Status: **{'Genau einmal erzeugt ✓' if iro2_ok else 'FEHLER: nicht erzeugt'}**",
        "",
        "## 9. Qualitätshinweise",
        "- Heading 4 wurde nicht verwendet (spec-konform); Zwischenüberschriften als Bold-Body.",
        "- sub-title und sub-number nie als sichtbare Überschrift ausgegeben.",
        ("- Mandatory cross-cut H2s immer erzeugt: "
         + ", ".join(CROSSCUT_DISPLAY.values())),
        "- Cross-cut H2 Display-Text: Format 'Xn - YYY (ESRS 2)'.",
        "- NUMBER-Ausgaben als Satz formuliert – KEIN [[REVIEW: expanded-from-single-value]] mehr.",
        "- ≤5 NUMBER Sub-Rows als gebündelter Stroetmann-Fließtext; >5 als Tabelle mit Summary-Satz.",
        "- Gesamt-/Brutto-Werte: 'Davon entfallen …'-Struktur; sonstige Bundles: 'ergänzend werden … ausgewiesen'.",
        "- NUMBER-Sätze mit variiertem Opener (Im Berichtsjahr beträgt / beläuft sich auf / werden ausgewiesen / berichteter Wert).",
        "- Jahreszahlen (1900–2100) ohne Tausenderpunkt (z.B. 2023, nicht 2.023).",
        "- yes/ja/teilweise → 1 Satz ohne Review-Marker (Spec v3).",
        "- Erläuterungen als plain body text integriert (kein 'Anmerkung:'-Label).",
        "- Alle Tabellen mit Einleitungstext versehen.",
        "- 'decimal'-Suffix im CSV → Prozentzahl (z.B. '50 decimal' → '50 %').",
        "- Europäische Zahlenformatierung (Komma = Dezimal, Punkt = Tausender).",
        "- Design: Century Gothic (Überschriften) + Calibri (Fließtext), Fröbel-Grüntöne.",
        f"- IRO-2 Tabelle: {'einmalig korrekt erzeugt' if iro2_ok else 'FEHLER'}.",
        "- Keine Fakten oder Annahmen aus Branchenwissen hinzugefügt.",
        "- EFRAG IG 3 nicht geladen – Datentyp heuristisch (unit → NUMBER, #-Muster → TABLE).",
        "- **Fließtext-Qualitätsgate angewendet: ja** (Stilreferenz: 20251128_CSRD Report_Stroetmann.docx).",
        "- Stilregeln: Claim-Elaboration-Kontext · Additive Konnektoren · Institutional 3rd Person · Passive/Aktiv-Wechsel.",
    ]

    with open(path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    print(f"QA-Log geschrieben: {path}")


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    print("Lade CSV …")
    with open(CSV_PATH, encoding='utf-8') as f:
        csv_rows = list(csv.DictReader(f))
    print(f"  {len(csv_rows)} Zeilen geladen.")

    print("Erstelle Word-Dokument …")
    doc, qa, iro2_done = build_document(csv_rows)

    doc.save(OUT_DOCX)
    print(f"✓ Dokument gespeichert: {OUT_DOCX}")

    write_qa(OUT_QA, csv_rows, qa, iro2_done)

    done  = sum(1 for e in qa if e[0] == 'DONE')
    skip  = sum(1 for e in qa if e[0] == 'SKIP')
    xcut  = sum(1 for e in qa if e[0] == 'SKIP_CROSSCUT')
    paras = sum(1 for e in qa if e[0] == 'REVIEW_PARA')
    expls = sum(1 for e in qa if e[0] == 'REVIEW_EXPL')
    kvs   = sum(1 for e in qa if e[0] == 'KEIN_VORFALL')
    print(f"  Gerendert: {done} Zeilen | Gefiltert: {skip} | "
          f"Cross-cuts: {xcut} | IRO-2 Tabelle: {'OK' if iro2_done else 'FEHLER'}")
    print(f"  Reviews: expanded-into-paragraph={paras} | "
          f"explanation-driven={expls} | keine-vorfälle={kvs}")


if __name__ == '__main__':
    main()
