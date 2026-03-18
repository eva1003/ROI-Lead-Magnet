#!/usr/bin/env python3
"""Transform CSRD_Report_v7.1 → v7.2: Convert remaining list-style Normal paragraphs
in the E1 (Umwelt) section into flowing German prose paragraphs."""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

INPUT  = 'Fröbel 2024/CSRD_Report_v7.1.docx'
OUTPUT = 'Fröbel 2024/CSRD_Report_v7.2.docx'

doc = Document(INPUT)
paras = doc.paragraphs  # live list – do NOT cache length


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_para_text(para, text):
    """Replace all runs in a paragraph with a single run containing text."""
    p_elem = para._element
    for r in list(p_elem.findall(qn('w:r'))):
        p_elem.remove(r)
    for hl in list(p_elem.findall(qn('w:hyperlink'))):
        p_elem.remove(hl)
    new_r = OxmlElement('w:r')
    # Copy rPr from first existing run if available (to preserve font/size)
    new_t = OxmlElement('w:t')
    new_t.text = text
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    p_elem.append(new_r)


def del_para_by_index(idx):
    """Delete paragraph at index idx (from doc.paragraphs)."""
    p = doc.paragraphs[idx]
    p._element.getparent().remove(p._element)


def find_idx(fragment, start=0):
    """Return index of first paragraph containing fragment, starting at start."""
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if fragment in p.text:
            return i
    return None


# ══════════════════════════════════════════════════════════════════════════════
# All replacements are defined as (search_fragment, new_text, n_items_to_delete_after)
# Processed in reverse order of appearance so deletions don't shift earlier indices.
# ══════════════════════════════════════════════════════════════════════════════

# We collect (idx, new_text, n_delete_after) tuples, then sort descending by idx.
ops = []

# ── 1. Scope 1 header + 2 items → flowing Scope-1 paragraph ──────────────────
ops.append((
    'Scope 1 – Direkte Emissionen (Heizung & Fuhrpark)',
    'Im Bereich Scope 1 (direkte Emissionen aus Heizung und Fuhrpark) adressiert Fröbel '
    'die beiden wesentlichen Quellen. Für Heizungsanlagen auf Basis von Holzpellets, Heizöl '
    'und Erdgas sind energetische Sanierungen, Heizungsoptimierungen und die sukzessive '
    'Umstellung auf erneuerbare Wärmetechnologien vorgesehen; die Emissionen sollen von '
    '1.515 t CO₂e (2019) auf 900 t CO₂e (–40,6 %) bis 2030 und auf 600 t CO₂e (–60,4 %) '
    'bis 2040 sinken. Parallel dazu wird die Dienstfahrzeugflotte vollständig elektrifiziert – '
    '80 % Elektrofahrzeuge bis 2027 und 100 % bis 2030 –, was einer Reduktion von derzeit '
    '50,9 t CO₂e auf 0 t CO₂e entspricht.',
    2  # delete the 2 item paragraphs that follow
))

# ── 2. Scope 2 header + 2 items → flowing Scope-2 paragraph ──────────────────
ops.append((
    'Scope 2 – Indirekte Emissionen (Strom, Fernwärme)',
    'Im Scope-2-Bereich (indirekte Emissionen aus Strom und Fernwärme) hat Fröbel den '
    'zentralen Schritt bereits vollzogen: Seit 2024 wird in allen Einrichtungen mit eigenem '
    'Stromvertrag zu 100 % zertifizierter Ökostrom bezogen, flankiert durch den laufenden '
    'Ausbau von Photovoltaikanlagen auf derzeit 30 Standorten. Die Emissionen aus der '
    'Stromnutzung sollen so von 1.788 t CO₂e (2019) auf 100 t CO₂e (–94,4 %) bis 2030 '
    'sinken. Für den Fernwärmebereich werden klimafreundlichere Wärmemixe und Pilotprojekte '
    'für erneuerbare Wärmetechnologien angestrebt, um die Emissionen von 3.042 t CO₂e auf '
    '700 t CO₂e (–77 %) bis 2030 zu reduzieren.',
    2  # delete the 2 item paragraphs that follow
))

# ── 3. Scope 3 header + 6 items → flowing Scope-3 paragraph ─────────────────
ops.append((
    'Scope 3 – Weitere indirekte Emissionen',
    'Den größten Hebel bildet der Scope-3-Bereich (weitere indirekte Emissionen). Mit rund '
    '40 % der Gesamtemissionen ist die Kantine und Lebensmittelversorgung die bedeutendste '
    'Stellschraube: Über den Fröbel-Leitfaden für Küchenteams (bereits zu 85,7 % umgesetzt), '
    'die Einführung von DGE-Standards, die vollständige Umstellung auf ovo-lacto-vegetarische '
    'Ernährung bis 2030 sowie den verstärkten Einsatz regionaler und biologischer Produkte '
    'sollen die Emissionen von 6.685 t CO₂e (2019) auf 3.725 t CO₂e (–44,3 %) bis 2030 '
    'sinken. Weitere Maßnahmen adressieren die Pendelemissionen der Mitarbeitenden durch das '
    'Deutschlandticket Job und digitale Arbeitsformate (Ziel: –4,5 % bis 2030), die '
    'Emissionen aus dem Familienverkehr durch Sensibilisierungskampagnen (Ziel: –27,3 % '
    'bis 2030), Geschäftsreisen durch die Reisekostenrichtlinie mit Bahnvorrang (Ziel: '
    '–60,7 % bis 2030) sowie Abfall und Recycling durch verbesserte Trennung und '
    'Lebensmittelabfallvermeidung. Ergänzend tragen Effizienzmaßnahmen bei Elektronik, '
    'Büromaterial und Wasser zur Gesamtreduktion bei.',
    6  # delete the 6 item paragraphs that follow
))

# ── 4. Bestandsgebäude intro paragraph (trailing colon → sentence) ────────────
ops.append((
    'Fröbel betreibt keine energie- oder rohstoffintensiven Anlagen',
    'Fröbel betreibt keine energie- oder rohstoffintensiven Anlagen und ist daher nur in '
    'geringem Umfang von gebundenen Treibhausgasemissionen betroffen. Die wesentlichen '
    'potenziellen Emissionsbindungen ergeben sich aus zwei strukturellen Risikofaktoren.',
    0
))

# ── 5. Bestandsgebäude block: intro + 3 items → one paragraph ────────────────
ops.append((
    'Ein wesentlicher Risikofaktor sind Bestandsgebäude mit konventionellen Heizsystemen',
    'Ein wesentlicher Risikofaktor sind Bestandsgebäude mit konventionellen Heizsystemen '
    'auf Basis von Gas, Öl oder Fernwärme. Ein Teil der Fröbel-Einrichtungen befindet sich '
    'in angemieteten oder denkmalgeschützten Gebäuden, in denen eine kurzfristige Umstellung '
    'auf erneuerbare Wärmequellen nicht möglich ist. Diese Infrastruktur verursacht derzeit '
    'rund 1,5 t CO₂e (Scope 1), deren Reduktion nur im Rahmen von Eigentümerentscheidungen '
    'oder Sanierungszyklen realisierbar ist, was das Risiko birgt, dass die '
    'Dekarbonisierungsgeschwindigkeit durch bauliche Rahmenbedingungen gebremst wird.',
    3  # delete the 3 item paragraphs that follow
))

# ── 6. Verpflegung block: intro + 4 items → one paragraph ────────────────────
ops.append((
    'Auch das Verpflegungssystem und Lebensmittelversorgung stellt einen strukturellen Risikofaktor dar',
    'Auch das Verpflegungssystem stellt einen strukturellen Risikofaktor dar. Ein Teil der '
    'Emissionen bleibt langfristig an die Bereitstellung von Mahlzeiten gebunden; sie liegen '
    'überwiegend im Scope-3-Bereich (Wareneinsatz Lebensmittel bei Einrichtungen mit '
    'Frischküche sowie externes Catering: ca. 6,7 t CO₂e im Jahr 2019). Trotz umfangreicher '
    'Umstellungen auf ovo-lacto-vegetarische Ernährung, regional-saisonale Zutaten und '
    'DGE-Standards bleiben unvermeidbare Restemissionen bestehen, die voraussichtlich durch '
    'Kompensationsmaßnahmen ausgeglichen werden müssen.',
    4  # delete the 4 item paragraphs that follow
))

# ── 7. Fortschritte Energie + Gebäude ────────────────────────────────────────
ops.append((
    'Im Bereich Energie und Gebäude erzielte Fröbel folgende Fortschritte',
    'Im Bereich Energie und Gebäude erzielte Fröbel bis Ende 2024 wesentliche Fortschritte. '
    'Die Umstellung auf 100 % zertifizierten Ökostrom wurde in allen Einrichtungen mit '
    'eigenem Stromvertrag abgeschlossen. 30 Einrichtungen verfügen bereits über '
    'Photovoltaikanlagen; weitere Projekte befinden sich in Planung. Darüber hinaus wurden '
    'erste Energieberatungen und Pilotprojekte zur Heizungsoptimierung erfolgreich '
    'durchgeführt, und die Einführung eines datenbasierten Ressourcenmonitorings wurde '
    'vorbereitet.',
    4  # delete the 4 item paragraphs that follow
))

# ── 8. Fortschritte Ernährung ─────────────────────────────────────────────────
ops.append((
    'Im Bereich Ernährung wurden folgende Fortschritte erzielt',
    'Im Bereich Ernährung wurden bedeutende Schritte unternommen. Der Fröbel-Leitfaden für '
    'Küchenteams ist zu 85,7 % umgesetzt; die vegetarische Ernährung wurde im Bereich West '
    'vollständig eingeführt und wird schrittweise in weiteren Regionen ausgeweitet. Für '
    'Einrichtungen mit Frischküche startet 2025 der DGE-Zertifizierungsprozess nach den '
    'Standards der Deutschen Gesellschaft für Ernährung.',
    3  # delete the 3 item paragraphs that follow
))

# ── 9. Fortschritte Mobilität ─────────────────────────────────────────────────
ops.append((
    'Im Bereich Mobilität erzielte Fröbel folgende Ergebnisse',
    'Im Bereich Mobilität wurde das nachhaltige Mobilitätskonzept vollständig umgesetzt. '
    'Der Fuhrpark ist zu 54,9 % elektrifiziert (Stand Ende 2024); bis 2030 soll eine '
    'vollständige Umstellung auf Elektrofahrzeuge erfolgen. Das Deutschlandticket Job wurde '
    'eingeführt, und alternative Mobilitätsformen werden aktiv gefördert.',
    3  # delete the 3 item paragraphs that follow
))

# ── 10. Fortschritte Monitoring ───────────────────────────────────────────────
ops.append((
    'Im Bereich Monitoring und Steuerung wurden folgende Maßnahmen umgesetzt',
    'Im Bereich Monitoring und Steuerung wurden die Treibhausgasbilanzen für 2019, 2022, '
    '2023 und 2024 erstellt. Die Erfassung der Gebäudedaten wurde in einer FM-Datenbank '
    'weitgehend zentralisiert, die Integration von Nachhaltigkeitskennzahlen in das interne '
    'Kontrollsystem (IKS) wurde begonnen, und die jährliche Überprüfung der Fortschritte '
    'durch den Stab Nachhaltigkeit ist als fester Prozess etabliert.',
    4  # delete the 4 item paragraphs that follow
))

# ── 11. Richtlinien block: intro + 7 items → flowing prose ───────────────────
ops.append((
    'Fröbel verfügt über eine Reihe von Richtlinien, die die Umsetzung der Klimaziele steuern',
    'Fröbel steuert die Umsetzung seiner Klimaziele über ein Bündel verbindlicher Richtlinien '
    'und Standards, die für alle Einrichtungen gelten. Die übergeordnete Fröbel '
    'Unternehmensstrategie 2030 definiert die strategischen Klimaziele (–50 % Emissionen '
    'bis 2030, Netto-Null bis 2040); die Fröbel Nachhaltigkeitsstrategie – erstmals im Rahmen '
    'des DNK-Standards veröffentlicht – übersetzt diese in konkrete Handlungsfelder, Ziele '
    'und Maßnahmen. Für das Reise- und Mobilitätsverhalten gilt die Fröbel '
    'Reisekostenrichtlinie, die Mitarbeitende verpflichtet, Dienstreisen nach '
    'Nachhaltigkeitskriterien durchzuführen und Bahnreisen gegenüber Flugreisen sowie '
    'digitale Meetings zu bevorzugen. Die Fröbel Beschaffungsrichtlinie verankert '
    'Nachhaltigkeitsaspekte im gesamten Beschaffungsprozess; eine weitergehende '
    'Berücksichtigung von Klimaeffekten entlang der Lieferkette ist geplant. Der Fröbel '
    'Baustandard legt Anforderungen an nachhaltiges Planen und Bauen für Neubauten und '
    'Sanierungen fest. Das Fröbel Mobilitätskonzept rahmt die Elektrifizierung des Fuhrparks '
    'und die Förderung nachhaltiger Mobilitätsalternativen. Der Fröbel Leitfaden für '
    'Küchenteams schließlich definiert verbindliche Mindeststandards für eine '
    'klimafreundliche und gesunde Ernährung in allen Einrichtungen.',
    7  # delete the 7 item paragraphs that follow
))

# ── 12. Kohärenz mechanisms block ─────────────────────────────────────────────
ops.append((
    'Die Kohärenz zwischen Zielsystem und Inventar wird durch folgende Mechanismen sichergestellt',
    'Die Kohärenz zwischen Zielsystem und Inventar stellt Fröbel durch mehrere Mechanismen '
    'sicher. Zieldefinition und THG-Bilanz folgen denselben organisatorischen und operativen '
    'Systemgrenzen (Konsolidierung nach operativer Kontrolle über alle Einrichtungen der '
    'Fröbel Bildung und Erziehung gGmbH); alle Berechnungen erfolgen nach einheitlicher '
    'Methodik gemäß GHG Protocol unter Verwendung des Planted-Tools. Die Ziele sind als '
    'Bruttoreduktionsziele formuliert und beziehen sich ausschließlich auf reale '
    'Emissionsminderungen innerhalb der eigenen Wertschöpfungskette – Kompensationen und '
    'CO₂-Zertifikate werden nicht eingerechnet. Die im Transitionsplan festgelegten '
    'Handlungsfelder (Gebäude, Energie, Ernährung, Mobilität) sind quantitativ mit den '
    'Hauptemissionsquellen im Inventar verknüpft, sodass Fortschritte messbar sind; die '
    'THG-Bilanzen werden jährlich aktualisiert, um die Zielerreichung zu überprüfen und '
    'Abweichungen frühzeitig zu identifizieren.',
    5  # delete the 5 mechanism item paragraphs that follow
))

# ── 13. Maßnahmen: Energie & Gebäude section ─────────────────────────────────
ops.append((
    '1. Energie und Gebäude (Scope 1 + 2):',
    'Im Bereich Energie und Gebäude (Scope 1 + 2) führt Fröbel Energieaudits durch und setzt '
    'Effizienzmaßnahmen wie Dämmung sowie Fenster- und Anlagentausch um (CapEx ≈ 3 Mio. € '
    'über das Globalbudget 2024–2025). Fossile Heizsysteme werden sukzessive durch '
    'Wärmepumpen oder Hybridlösungen ersetzt. Der Ausbau der Photovoltaik auf Gebäuden in '
    'Eigentum und Erbbaurecht schreitet voran (aktuell 30 Einrichtungen); der vollständige '
    'Strombezug über zertifizierten Ökostrom wurde bereits abgeschlossen. Bis 2026 ist die '
    'Einführung eines datenbasierten Energiemanagementsystems geplant.',
    5  # delete the 5 item paragraphs that follow
))

# ── 14. Maßnahmen: Ernährung section ─────────────────────────────────────────
ops.append((
    '2. Ernährung und Beschaffung (Scope 3):',
    'Im Handlungsfeld Ernährung und Beschaffung (Scope 3) steht die Umsetzung des '
    'Fröbel-Leitfadens für Küchenteams im Mittelpunkt, der bereits zu 85,7 % eingeführt '
    'ist. Speisepläne werden stufenweise auf ovo-lacto-vegetarische Kost umgestellt (Bereich '
    'West abgeschlossen, übrige Regionen bis 2030); verbindliche DGE-Standards für alle '
    'Küchen starten 2025. Durch den verstärkten Einsatz ökologisch zertifizierter, regionaler '
    'und saisonaler Produkte sowie die Reduktion tierischer Lebensmittel werden die '
    'ernährungsbedingten Scope-3-Emissionen gezielt gesenkt.',
    4  # delete the 4 item paragraphs that follow
))

# ── 15. Maßnahmen: Mobilität section ─────────────────────────────────────────
ops.append((
    '3. Mobilität (Scope 1 + 3):',
    'Im Bereich Mobilität (Scope 1 + 3) wird der Fuhrpark konsequent elektrifiziert: '
    '80 % Elektrofahrzeuge sind bis 2027 angestrebt, 100 % bis 2030. Das Deutschlandticket '
    'Job ist eingeführt, JobRad-Angebote werden ausgebaut, und ÖPNV-Nutzung wird durch '
    'gezielte Anreize gefördert. Die Reisekostenrichtlinie verpflichtet zudem zur Nutzung '
    'nachhaltiger Reiseoptionen – mit Priorität für Bahnreisen gegenüber Flugreisen und dem '
    'Vorrang digitaler Meetings.',
    3  # delete the 3 item paragraphs that follow
))

# ── 16. Maßnahmen: Monitoring section ────────────────────────────────────────
ops.append((
    '4. Monitoring und Steuerung:',
    'Das vierte Handlungsfeld – Monitoring und Steuerung – sichert die Wirksamkeit aller '
    'Maßnahmen ab. Treibhausgasbilanzen für die Jahre 2019, 2022–2024 liegen vor und werden '
    'künftig jährlich erstellt. Eine automatisierte Datenerfassung in den Bereichen Energie, '
    'Fuhrpark und Beschaffung ist im Aufbau; die Integration der Nachhaltigkeitskennzahlen '
    'in das interne Kontrollsystem (IKS) ermöglicht eine laufende Evaluierung der '
    'Fortschritte gegenüber den Reduktionszielen.',
    3  # delete the 3 item paragraphs that follow
))


# ══════════════════════════════════════════════════════════════════════════════
# Execute operations in reverse document order so index shifts don't matter
# ══════════════════════════════════════════════════════════════════════════════

# Resolve fragment → index (forward pass first)
resolved = []
for fragment, new_text, n_del in ops:
    idx = find_idx(fragment)
    if idx is None:
        print(f'⚠  Not found: {repr(fragment[:60])}')
    else:
        resolved.append((idx, new_text, n_del, fragment))

# Sort descending so deletions don't shift earlier paragraph indices
resolved.sort(key=lambda x: x[0], reverse=True)

for idx, new_text, n_del, fragment in resolved:
    # Delete following item paragraphs first (highest index first)
    for i in range(idx + n_del, idx, -1):
        del_para_by_index(i)
    # Replace intro paragraph text
    set_para_text(doc.paragraphs[idx], new_text)
    print(f'✓ idx {idx:3d}: replaced + deleted {n_del} items  [{fragment[:50]}]')


# ══════════════════════════════════════════════════════════════════════════════
# Verify: count remaining list-style paragraphs in Umwelt / E1 section
# ══════════════════════════════════════════════════════════════════════════════

e1_start = next((i for i, p in enumerate(doc.paragraphs)
                 if 'E1-1' in p.text and p.style.name.startswith('Heading')), 0)
e2_start = next((i for i, p in enumerate(doc.paragraphs)
                 if 'E2 - IRO-1' in p.text and p.style.name.startswith('Heading')),
                len(doc.paragraphs))

# Check for any remaining "label: content" sentence-initial patterns typical of list items
remaining = []
for p in doc.paragraphs[e1_start:e2_start]:
    txt = p.text.strip()
    if (p.style.name == 'Normal' and txt and
            not txt.startswith('Die ') and not txt.startswith('Der ') and
            not txt.startswith('Fröbel') and not txt.startswith('Im ') and
            not txt.startswith('Für ') and not txt.startswith('Alle ') and
            not txt.startswith('Zur ') and not txt.startswith('Damit') and
            not txt.startswith('Das ') and not txt.startswith('Wo ') and
            not txt.startswith('Ein ') and not txt.startswith('Auch ') and
            not txt.startswith('Trotz') and not txt.startswith('Es ') and
            not txt.startswith('Damit') and not txt.startswith('In ') and
            not txt.startswith('Auf ') and not txt.startswith('Sollten') and
            not txt.startswith('Bei ') and not txt.startswith('ESRS ') and
            not txt.startswith('Methodi') and not txt.startswith('Erläut') and
            not txt.startswith('Ergänz') and not txt.startswith('Informati')):
        remaining.append(txt[:80])

print(f'\n--- Potentially remaining list-style paragraphs in E1 ({len(remaining)}) ---')
for r in remaining:
    print(f'  {r}')

doc.save(OUTPUT)
print(f'\n✓ Gespeichert: {OUTPUT}')
