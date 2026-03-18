#!/usr/bin/env python3
"""Remove all bullet points from E1 (Umwelt) section – convert to flowing text."""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = 'Fröbel 2024/CSRD_Report_v7.1.docx'
OUTPUT = 'Fröbel 2024/CSRD_Report_v7.1.docx'

doc = Document(INPUT)

# ── Helpers ────────────────────────────────────────────────────────────────────

def find_para(fragment, skip=0):
    seen = 0
    for p in doc.paragraphs:
        if fragment in p.text:
            if seen == skip:
                return p
            seen += 1
    return None

def del_para(para):
    para._element.getparent().remove(para._element)

def set_para_text(para, text):
    p_elem = para._element
    for r in list(p_elem.findall(qn('w:r'))):
        p_elem.remove(r)
    for hl in list(p_elem.findall(qn('w:hyperlink'))):
        p_elem.remove(hl)
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    if text.startswith(' ') or text.endswith(' '):
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    p_elem.append(new_r)

def get_list_block_after(intro_para):
    """Return list of consecutive List Bullet paragraphs immediately after intro_para."""
    items = []
    e = intro_para._element
    nxt = e.getnext()
    while nxt is not None and nxt.tag.endswith('}p'):
        pPr = nxt.find(qn('w:pPr'))
        style_val = ''
        if pPr is not None:
            pStyle = pPr.find(qn('w:pStyle'))
            if pStyle is not None:
                style_val = pStyle.get(qn('w:val')) or ''
        if 'List' in style_val or 'list' in style_val:
            # Reconstruct text from w:t nodes
            t_nodes = nxt.findall('.//' + qn('w:t'))
            text = ''.join(t.text or '' for t in t_nodes)
            items.append((nxt, text))
            nxt = nxt.getnext()
        else:
            break
    return items

def merge_items_into_intro(intro_para, separator=' ', strip_intro_colon=True):
    """Merge all following List Bullet items into intro_para as flowing text, delete items."""
    items = get_list_block_after(intro_para)
    if not items:
        return
    intro_text = intro_para.text.rstrip()
    if strip_intro_colon and intro_text.endswith(':'):
        intro_text = intro_text[:-1]
    item_texts = [text.strip().rstrip(',') for (_, text) in items]
    merged = intro_text + separator + ', '.join(item_texts[:-1])
    if len(item_texts) > 1:
        merged += ' sowie ' + item_texts[-1]
    else:
        merged += ' ' + item_texts[0]
    if not merged.endswith('.'):
        merged += '.'
    set_para_text(intro_para, merged)
    for (elem, _) in items:
        elem.getparent().remove(elem)

def convert_block_to_normal(intro_para, remove_intro_colon=False):
    """Change style of all following List Bullet items to Normal."""
    items_elems = get_list_block_after(intro_para)
    if not items_elems:
        return 0
    if remove_intro_colon:
        t = intro_para.text.rstrip()
        if t.endswith(':'):
            set_para_text(intro_para, t[:-1])
    for (elem, _) in items_elems:
        pPr = elem.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            elem.insert(0, pPr)
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            pStyle = OxmlElement('w:pStyle')
            pPr.insert(0, pStyle)
        pStyle.set(qn('w:val'), 'Normal')
        # Also remove indentation set by list style
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)
    return len(items_elems)


# ══════════════════════════════════════════════════════════════════════════════
# PHASE 1: Explicit merges for grammatically incomplete item lists
# ══════════════════════════════════════════════════════════════════════════════

# ── 1a. Globalbudget: 4 dative-clause items → one sentence ───────────────────
p_budget = find_para('Für das Jahr 2024 wurde ein Globalbudget in Höhe von 1 Mio. Euro beschlossen. Die Mittel dienen insbesondere:')
if p_budget:
    items = get_list_block_after(p_budget)
    item_texts = [t.strip().rstrip(',').rstrip('.') for (_, t) in items]
    # Build proper German sentence: "insbesondere X, Y, Z sowie W."
    base = 'Für das Jahr 2024 wurde ein Globalbudget in Höhe von 1 Mio. Euro beschlossen. Die Mittel dienen insbesondere '
    if len(item_texts) >= 2:
        merged = base + ', '.join(item_texts[:-1]) + ' sowie ' + item_texts[-1] + '.'
    else:
        merged = base + ' '.join(item_texts) + '.'
    set_para_text(p_budget, merged)
    for (elem, _) in items:
        elem.getparent().remove(elem)
    print(f"✓ Globalbudget: merged {len(items)} items into prose")

# ── 1b. Strategy links: 3 "Gebäude / Ernährung / Mobilität" items ─────────────
p_strat = find_para('direkt mit den Handlungsfeldern der Geschäftsstrategie verknüpft – insbesondere:')
if p_strat:
    items = get_list_block_after(p_strat)
    # Build: "...insbesondere mit Gebäude & Energie (...), Ernährung & Verpflegung (...) sowie Mobilität (...)"
    def to_bracket(text):
        """Convert 'Label: description,' → 'Label (description)'"""
        text = text.strip().rstrip(',').rstrip('.')
        if ': ' in text:
            label, desc = text.split(': ', 1)
            return f'{label} ({desc})'
        return text
    parts = [to_bracket(t) for (_, t) in items]
    base = 'Die Maßnahmen des Transitionsplans sind direkt mit den Handlungsfeldern der Geschäftsstrategie verknüpft – insbesondere mit '
    if len(parts) >= 2:
        merged = base + ', '.join(parts[:-1]) + ' sowie ' + parts[-1] + '.'
    else:
        merged = base + parts[0] + '.'
    set_para_text(p_strat, merged)
    for (elem, _) in items:
        elem.getparent().remove(elem)
    print(f"✓ Strategy links: merged {len(items)} items into prose")

# ── 1c. würde Fröbel: 3 verb-phrase items → one sentence ─────────────────────
p_waere = find_para('In einem solchen Fall würde Fröbel:')
if p_waere:
    items = get_list_block_after(p_waere)
    item_texts = [t.strip().rstrip(',') for (_, t) in items]
    # "würde Fröbel X, Y und Z."
    if len(item_texts) >= 2:
        merged = ('In einem solchen Fall würde Fröbel ' +
                  ', '.join(item_texts[:-1]) + ' und ' + item_texts[-1])
    else:
        merged = 'In einem solchen Fall würde Fröbel ' + item_texts[0]
    if not merged.endswith('.'):
        merged += '.'
    set_para_text(p_waere, merged)
    for (elem, _) in items:
        elem.getparent().remove(elem)
    print(f"✓ würde Fröbel: merged {len(items)} items into one sentence")


# ══════════════════════════════════════════════════════════════════════════════
# PHASE 2: Convert remaining List Bullet blocks to Normal paragraphs
# ══════════════════════════════════════════════════════════════════════════════

# Scope 1 (2 long items)
p = find_para('Scope 1 – Direkte Emissionen (Heizung & Fuhrpark):')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=True)
    print(f"✓ Scope 1 block: {n} items → Normal")

# Scope 2 (2 long items)
p = find_para('Scope 2 – Indirekte Emissionen (Strom, Fernwärme):')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=True)
    print(f"✓ Scope 2 block: {n} items → Normal")

# Scope 3 (6 long items)
p = find_para('Scope 3 – Weitere indirekte Emissionen:')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=True)
    print(f"✓ Scope 3 block: {n} items → Normal")

# Bestandsgebäude (3 items – complete sentences)
p = find_para('bestandsgebäuden mit konventionellen heizsystemen')
if p:
    # Also fix lowercase intro
    set_para_text(p, 'Ein wesentlicher Risikofaktor sind Bestandsgebäude mit konventionellen Heizsystemen (Gas, Öl, Fernwärme):')
    n = convert_block_to_normal(p, remove_intro_colon=False)
    print(f"✓ Bestandsgebäude block: {n} items → Normal, intro capitalized")

# Verpflegung (4 items – complete sentences)
p = find_para('Verpflegungssystem und Lebensmittelversorgung stellt einen strukturellen Risikofaktor dar:')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=False)
    print(f"✓ Verpflegung block: {n} items → Normal")

# Progress: Energie & Gebäude (4 items)
p = find_para('Im Bereich Energie und Gebäude erzielte Fröbel folgende Fortschritte:')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=False)
    print(f"✓ Progress-Energie block: {n} items → Normal")

# Progress: Ernährung (3 items)
p = find_para('Im Bereich Ernährung wurden folgende Fortschritte erzielt:')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=False)
    print(f"✓ Progress-Ernährung block: {n} items → Normal")

# Progress: Mobilität (3 items)
p = find_para('Im Bereich Mobilität erzielte Fröbel folgende Ergebnisse:')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=False)
    print(f"✓ Progress-Mobilität block: {n} items → Normal")

# Progress: Monitoring (4 items)
p = find_para('Im Bereich Monitoring und Steuerung wurden folgende Maßnahmen umgesetzt:')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=False)
    print(f"✓ Progress-Monitoring block: {n} items → Normal")

# Richtlinien (8 items: "Fröbel Policy: description")
p = find_para('Fröbel verfügt über eine Reihe von Richtlinien, die die Umsetzung der Klimaziele steuern')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=False)
    print(f"✓ Richtlinien block: {n} items → Normal")

# Kohärenz (5 items: "Concept: description")
p = find_para('Kohärenz zwischen Zielsystem und Inventar wird durch folgende Mechanismen sichergestellt:')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=False)
    print(f"✓ Kohärenz block: {n} items → Normal")

# Action block 1: Energie & Gebäude
p = find_para('1. Energie und Gebäude (Scope 1 + 2):')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=False)
    print(f"✓ Action-Energie block: {n} items → Normal")

# Action block 2: Ernährung
p = find_para('2. Ernährung und Beschaffung (Scope 3):')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=False)
    print(f"✓ Action-Ernährung block: {n} items → Normal")

# Action block 3: Mobilität
p = find_para('3. Mobilität (Scope 1 + 3):')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=False)
    print(f"✓ Action-Mobilität block: {n} items → Normal")

# Action block 4: Monitoring
p = find_para('4. Monitoring und Steuerung:')
if p:
    n = convert_block_to_normal(p, remove_intro_colon=False)
    print(f"✓ Action-Monitoring block: {n} items → Normal")


# ══════════════════════════════════════════════════════════════════════════════
# PHASE 3: Safety net – convert any remaining List Bullet in E1 section
# ══════════════════════════════════════════════════════════════════════════════

e1_start = next((i for i, p in enumerate(doc.paragraphs)
                 if 'E1-1' in p.text and p.style.name.startswith('Heading')), 0)
e2_start = next((i for i, p in enumerate(doc.paragraphs)
                 if 'E2 - IRO-1' in p.text and p.style.name.startswith('Heading')),
                len(doc.paragraphs))

remaining = 0
for p in doc.paragraphs[e1_start:e2_start]:
    if p.style.name == 'List Bullet':
        p.style = doc.styles['Normal']
        remaining += 1

if remaining:
    print(f"✓ Safety net: converted {remaining} remaining List Bullet → Normal in E1")


# ══════════════════════════════════════════════════════════════════════════════
# PHASE 4: Verify – count residual bullets in E1
# ══════════════════════════════════════════════════════════════════════════════

still_bullets = sum(1 for p in doc.paragraphs[e1_start:e2_start]
                    if p.style.name == 'List Bullet')
print(f"\n{'✓' if still_bullets == 0 else '❌'} Residual List Bullet in E1: {still_bullets}")

doc.save(OUTPUT)
print(f"✓ Gespeichert: {OUTPUT}")
