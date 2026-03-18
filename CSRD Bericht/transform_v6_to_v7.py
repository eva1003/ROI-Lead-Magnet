#!/usr/bin/env python3
"""
CSRD_Report_v6 → CSRD_Report_v7
Allgemeine Informationen (ESRS 2) — Textredaktion

Änderungen:
- Einzelworte zu Sätzen
- Duplikate entfernt
- Spiegelstriche durch Fließtext / schlüssige Aufzählungen ersetzt
- Lesefluss verbessert (Stroetmann-Stil)
"""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = 'Fröbel 2024/CSRD_Report_v6.docx'
OUTPUT = 'Fröbel 2024/CSRD_Report_v7.docx'

doc = Document(INPUT)

# ─── helpers ───────────────────────────────────────────────────────────────

def find_para(fragment, skip=0):
    """Return paragraph containing fragment (skip=n skips first n matches)."""
    seen = 0
    for p in doc.paragraphs:
        if fragment in p.text:
            if seen == skip:
                return p
            seen += 1
    return None

def find_all(fragment):
    return [p for p in doc.paragraphs if fragment in p.text]

def set_text(para, text):
    """Replace paragraph content, preserving first-run character formatting."""
    for i, run in enumerate(para.runs):
        run.text = text if i == 0 else ''
    if not para.runs:
        para.add_run(text)

def del_para(para):
    e = para._element
    e.getparent().remove(e)

def _list_style():
    available = {s.name for s in doc.styles}
    for name in ('List Bullet', 'List Paragraph', 'Normal'):
        if name in available:
            return name
    return 'Normal'

LIST_STYLE = _list_style()

def add_list_items_after(ref_para, items):
    """Insert properly styled list-item paragraphs after ref_para (in order)."""
    prev = ref_para
    for item in items:
        new_p = doc.add_paragraph()
        new_p.style = doc.styles[LIST_STYLE]
        new_p.add_run(item)
        prev._element.addnext(new_p._element)
        prev = new_p
    return prev

def split_numbered_para(para, intro_text):
    """
    Split 'intro_text' + '1. Step A. 2. Step B. ...' paragraph
    into intro paragraph + numbered items.
    Returns list of item texts.
    """
    import re
    full = para.text
    # extract numbered parts
    parts = re.split(r'\s*\d+\.\s+', full)
    items = [p.strip() for p in parts if p.strip()]
    set_text(para, intro_text)
    return items

# ═══════════════════════════════════════════════════════════════════════════
# BP-1 — Allgemeine Grundlagen
# ═══════════════════════════════════════════════════════════════════════════

# [15] "Konsolidiert"  →  delete; improve [16] opening
p15 = find_para('Konsolidiert')
if p15 and p15.text.strip() == 'Konsolidiert':
    del_para(p15)

p16 = find_para('Ja. Der Konsolidierungskreis')
if p16:
    set_text(p16,
        'Die Nachhaltigkeitserklärung wird auf konsolidierter Basis erstellt. '
        'Der Konsolidierungskreis entspricht dem des Konzernabschlusses und '
        'umfasst die vollkonsolidierten Tochtergesellschaften Fröbel Bildung '
        'und Erziehung gGmbH, FRÖBEL Kinder in Bewegung gGmbH und '
        'FRÖBEL Akademie gGmbH.'
    )

# ═══════════════════════════════════════════════════════════════════════════
# GOV-1 — Zusammensetzung der Organe
# ═══════════════════════════════════════════════════════════════════════════

# Collect the three repeated "berichtete Wert..." paragraphs BEFORE any deletion
val_paras = find_all('berichtete Wert für zusammensetzung')
# Also the "Zur Einordnung" duplicate
einordnung_para = find_para('Zur Einordnung ist hinzuzufügen')
# The "36. Gewählte Betriebsräte" sentence
betriebsraete_para = find_para('36. Gewählte Betriebsräte')

# Rewrite the first one (beträgt 5) → combined authoritative paragraph
if val_paras:
    p22 = val_paras[0]
    set_text(p22,
        'Die Verwaltungs-, Leitungs- und Aufsichtsorgane der FRÖBEL-Gruppe '
        'umfassen insgesamt 31 Mitglieder: 5 geschäftsführende und '
        '26 nicht geschäftsführende, was einem Anteil von 83,9\u202f% nicht '
        'geschäftsführender Mitglieder entspricht. Die Verteilung nach '
        'Organisationseinheiten stellt sich wie folgt dar '
        '(geschäftsführend\u202f/\u202fnicht geschäftsführend): '
        'Fröbel e.\u202fV. mit 1\u202f/\u202f1, Fröbel Bildung und Erziehung '
        'gGmbH mit 1\u202f/\u202f11, FRÖBEL Kinder in Bewegung gGmbH mit '
        '1\u202f/\u202f3, FRÖBEL Akademie gGmbH mit 2\u202f/\u202f4 sowie '
        'der Fröbel Aufsichtsrat mit 0\u202f/\u202f7. Stefan Spieker übt als '
        'Vorstandsvorsitzender des Fröbel e.\u202fV. sowie als Geschäftsführer '
        'der Fröbel Bildung und Erziehung gGmbH, der FRÖBEL Kinder in Bewegung '
        'gGmbH und der FRÖBEL Akademie gGmbH eine Mehrfachfunktion aus; bei der '
        'Zählung werden Rollen, nicht Personen erfasst. Im Jahr 2022 wurden '
        'darüber hinaus in den Regionen West, Mitte und Ost insgesamt 36 '
        'Betriebsräte gewählt, die zum Teil auch den Gesamtbetriebsrat sowie '
        'die Schwerbehindertenvertretungen stellen.'
    )

# Delete duplicates
for dup in val_paras[1:]:
    del_para(dup)
if betriebsraete_para:
    del_para(betriebsraete_para)
if einordnung_para:
    del_para(einordnung_para)

# Fix both table intros (same text → find both at once)
tbl_intros = find_all('Die folgende Tabelle zeigt die Angaben zu zusammensetzung')
if len(tbl_intros) >= 1:
    set_text(tbl_intros[0],
        'Die folgende Tabelle gibt einen Überblick über die '
        'Geschlechterverteilung in den Verwaltungs-, Leitungs- und '
        'Aufsichtsorganen.'
    )
if len(tbl_intros) >= 2:
    set_text(tbl_intros[1],
        'Die folgende Tabelle gibt einen Überblick über die '
        'Geschlechterverteilung der Beschäftigten.'
    )

# ═══════════════════════════════════════════════════════════════════════════
# GOV-2 — Befasste Themen: inline em-dash list → proper list
# ═══════════════════════════════════════════════════════════════════════════

p62 = find_para('befassten sich die Verwaltungs-, Leitungs- und/oder Aufsichtsorgane')
p63 = find_para('Energieverbrauch und Treibhausgasemissionen aus dem Betrieb von Gebäuden')
if p63 and p62:
    del_para(p63)
    add_list_items_after(p62, [
        'Energieverbrauch und Treibhausgasemissionen aus dem Betrieb von Gebäuden',
        'Klimaschutzmaßnahmen sowie die Reduktion klimabezogener Auswirkungen',
        'Klimabezogene Risiken für den Betrieb, insbesondere im Zusammenhang '
        'mit Hitze und Extremwetterereignissen',
        'Arbeitsbedingungen und Beschäftigungssituation der Mitarbeitenden',
        'Fachkräftesicherung und Personalgewinnung',
        'Qualität der frühkindlichen Bildung',
        'Bildung für nachhaltige Entwicklung (BNE)',
        'Gesellschaftliche Wirkung der pädagogischen Arbeit',
        'Weiterentwicklung der Nachhaltigkeitsstrategie',
        'Durchführung der doppelten Wesentlichkeitsanalyse',
        'Nachhaltigkeitsbezogene Risiken und Chancen',
        'Anforderungen aus der CSRD, den ESRS sowie der '
        'Nachhaltigkeitsberichterstattung',
    ])

# ═══════════════════════════════════════════════════════════════════════════
# GOV-5 — Risikomanagement: inline bullet lists → Fließtext
# ═══════════════════════════════════════════════════════════════════════════

# [83] "Zentrale Bestandteile sind: - ..."
p83 = find_para('Zentrale Bestandteile sind:')
if p83:
    set_text(p83,
        'Zentrale Bestandteile des Systems sind klar definierte '
        'Zuständigkeiten für die Erhebung, Prüfung und Freigabe '
        'nachhaltigkeitsbezogener Daten, standardisierte Erhebungs- und '
        'Dokumentationsprozesse für qualitative und quantitative Angaben '
        '(z.\u202fB. Energie, Mobilität, Verpflegung, Arbeitskräfte, '
        'Kinderschutz), fachliche Plausibilitätsprüfungen durch zuständige '
        'Bereiche sowie ein mehrstufiges internes Review-Verfahren auf '
        'operativer, fachlicher und zentraler Ebene.'
    )

# [86] "Risiken werden anhand folgender Kriterien bewertet und priorisiert: - ..."
p86 = find_para('Risiken werden anhand folgender Kriterien bewertet')
if p86:
    set_text(p86,
        'Risiken werden anhand folgender Kriterien bewertet und priorisiert: '
        'potenzielle Auswirkungen auf Menschen, Umwelt und Unternehmenssteuerung, '
        'Eintrittswahrscheinlichkeit von Fehldarstellungen oder unvollständigen '
        'Angaben, Datenverfügbarkeit und -qualität (Primär- vs. Sekundärdaten, '
        'Modellierungen) sowie regulatorische Relevanz und Prüfungsrisiken.'
    )

# [88] "Zu den wesentlichen Risiken ... zählen insbesondere: - ..."
p88 = find_para('wesentlichen Risiken in der Nachhaltigkeitsberichterstattung zählen')
if p88:
    set_text(p88,
        'Zu den wesentlichen Risiken in der Nachhaltigkeitsberichterstattung '
        'zählen insbesondere unvollständige oder uneinheitliche '
        'Datenverfügbarkeit über Standorte hinweg, die Abhängigkeit von '
        'modellierten Daten in einzelnen Themenfeldern (z.\u202fB. Mobilität, '
        'Verpflegung, vorgelagerte Emissionen), Schnittstellenrisiken zwischen '
        'Fachbereichen sowie Interpretationsspielräume bei neuen oder '
        'weiterentwickelten ESRS-Vorgaben.'
    )

# [89] "Zur Minderung dieser Risiken werden u. a. folgende Kontrollen angewendet: - ..."
p89 = find_para('Minderung dieser Risiken werden u.\u202fa.')
if not p89:
    p89 = find_para('Minderung dieser Risiken werden u. a.')
if p89:
    set_text(p89,
        'Zur Minderung dieser Risiken werden insbesondere folgende Kontrollen '
        'angewendet: die klare Definition von Datenquellen, Annahmen und '
        'Berechnungsmethoden, fachliche Prüfung und Freigabe durch zuständige '
        'Bereiche, Dokumentation zentraler Annahmen und Abgrenzungen sowie '
        'Konsistenzprüfungen zwischen Textangaben, Kennzahlen und '
        'Methodenbeschreibungen.'
    )

# [91] "Identifizierte Risiken und Verbesserungspotenziale werden genutzt zur: - ..."
p91 = find_para('Identifizierte Risiken und Verbesserungspotenziale werden genutzt zur')
if not p91:
    p91 = find_para('Identifizierte Risiken und Verbesserungspotenziale werden genutzt')
if p91:
    set_text(p91,
        'Identifizierte Risiken und Verbesserungspotenziale werden genutzt, um '
        'Erhebungs- und Dokumentationsprozesse anzupassen, Zuständigkeiten und '
        'Schnittstellen zu schärfen, interne Leitlinien und Arbeitshilfen '
        'weiterzuentwickeln sowie Maßnahmen zur Verbesserung der Datenqualität '
        'zu priorisieren.'
    )

# [94] "Dies erfolgt insbesondere: - ..."
p94 = find_para('Dies erfolgt insbesondere:')
if not p94:
    p94 = find_para('im Zuge der Erstellung und Abstimmung der Nachhaltigkeitserklärung')
if p94 and '- ' in p94.text:
    set_text(p94,
        'Dies erfolgt insbesondere im Zuge der Erstellung und Abstimmung der '
        'Nachhaltigkeitserklärung, im Rahmen von Berichten und Vorlagen an die '
        'Unternehmensleitung sowie bei Bedarf in Berichten an den Aufsichtsrat, '
        'insbesondere bei wesentlichen Änderungen, Risiken oder '
        'Weiterentwicklungen der Berichterstattung.'
    )

# ═══════════════════════════════════════════════════════════════════════════
# SBM-3 — Duplikate [115]–[118] entfernen (identisch mit [111]–[114])
# ═══════════════════════════════════════════════════════════════════════════

dup_fragments = [
    'wesentlichen Auswirkungen des Unternehmens ergeben sich im Schwerpunkt',
    'Im sozialen Bereich betreffen die wesentlichen Auswirkungen insbesondere Kinder',
    'Potenzielle negative Auswirkungen bestehen insbesondere im Zusammenhang mit Kinderschutz',
    'Im Umweltbereich ergeben sich wesentliche Auswirkungen aus dem Energie',
]

for frag in dup_fragments:
    occurrences = find_all(frag)
    # Keep the first, delete all subsequent
    for dup in occurrences[1:]:
        del_para(dup)

# ═══════════════════════════════════════════════════════════════════════════
# SBM-3 — [126] inline bullet list → Fließtext
# ═══════════════════════════════════════════════════════════════════════════

p126 = find_para('belegungsabhängigen Erlösen und Auslastungsrisiken')
if p126 and '- ' in p126.text:
    set_text(p126,
        'Aktuelle finanzielle Auswirkungen ergeben sich insbesondere aus '
        'belegungsabhängigen Erlösen und Auslastungsrisiken, aus Personal-, '
        'Betriebs- und Verpflegungskosten, aus Preisentwicklungen bei Energie '
        'und Lebensmitteln sowie aus potenziellen Kosten infolge von Qualitäts-, '
        'Hygiene- oder Kinderschutzvorfällen.'
    )

# ═══════════════════════════════════════════════════════════════════════════
# IRO-1 — [133] alle 5 Schritte in einem Absatz → separate Absätze
# ═══════════════════════════════════════════════════════════════════════════

p133 = find_para('Festlegung des Konsolidierungskreises und Abbildung')
if p133 and '2. Identifikation' in p133.text:
    import re
    raw = p133.text
    # Split on number-dot pattern
    parts = re.split(r'\s*(\d+\.\s)', raw)
    steps = []
    i = 1
    while i < len(parts):
        num = parts[i].strip()
        content = parts[i+1].strip() if i+1 < len(parts) else ''
        steps.append(num + ' ' + content)
        i += 2
    if steps:
        set_text(p133, steps[0])
        prev = p133
        for step in steps[1:]:
            new_p = doc.add_paragraph()
            new_p.style = doc.styles['Normal']
            new_p.add_run(step)
            prev._element.addnext(new_p._element)
            prev = new_p

# ═══════════════════════════════════════════════════════════════════════════
# IRO-1 — [136] "Das Vorgehen im Überblick: - ..." → Aufzählung
# ═══════════════════════════════════════════════════════════════════════════

p136 = find_para('Das Vorgehen im Überblick:')
if p136:
    set_text(p136, 'Das Vorgehen im Überblick:')
    add_list_items_after(p136, [
        'Identifikation der Auswirkungen entlang der Wertschöpfungskette '
        '(Scope 1–3), einschließlich Beschaffung von Waren, Lebensmitteln '
        'und Energie, Mobilität der Mitarbeitenden und Familien sowie '
        'Abfall und Recycling',
        'Bewertung der Auswirkungen anhand von Schweregrad, Umfang und '
        'Wahrscheinlichkeit (vgl. ESRS 1, Rz. 85\u202fff.)',
        'Priorisierung nach relativer Bedeutung für Mensch und Umwelt '
        '(vgl. ESRS 1, Rz. 87\u202fff.)',
        'Monitoring durch regelmäßige Überprüfung und Dokumentation im IKS '
        'sowie Integration in das allgemeine Nachhaltigkeitsmanagement',
    ])

# ═══════════════════════════════════════════════════════════════════════════
# IRO-1 — [146] Scoring-Kriterien als Aufzählung
# ═══════════════════════════════════════════════════════════════════════════

p146 = find_para('mehrdimensionales Scoring-System auf einer 5-Punkte-Skala eingesetzt')
if p146 and '- ' in p146.text:
    # Keep intro sentence, strip the inline list
    intro = p146.text.split(': -')[0].rstrip() + ':'
    set_text(p146, intro)
    add_list_items_after(p146, [
        'Eintrittswahrscheinlichkeit: von sehr unwahrscheinlich (1) bis '
        'sicherer Eintritt (5)',
        'Ausmaß positiver und negativer Auswirkungen: von minimal (1) bis '
        'sehr hoch (5)',
        'Umfang der Auswirkungen: von minimal (1) bis maximal/weitreichend (5)',
        'Wiederherstellbarkeit (nur negative Auswirkungen): von sehr einfach '
        '(1) bis unmöglich (5)',
    ])

# ═══════════════════════════════════════════════════════════════════════════
# IRO-1 — [150] Verfahrensschritte Finanzielle Wesentlichkeit → Aufzählung
# ═══════════════════════════════════════════════════════════════════════════

p150 = find_para('Die Analyse folgte einem mehrstufigen Vorgehen:')
if p150 and '- ' in p150.text:
    intro = p150.text.split(': -')[0].rstrip() + ':'
    set_text(p150, intro)
    add_list_items_after(p150, [
        'Identifikation: Erfassung potenzieller finanzieller Risiken und '
        'Chancen entlang der Wertschöpfungskette, insbesondere in Bezug auf '
        'Personalgewinnung und -bindung, Energieversorgung, '
        'Lebensmittelversorgung sowie externe Regulierungen',
        'Bewertung: Einordnung nach Eintrittswahrscheinlichkeit, Ausmaß und '
        'zeitlichem Horizont (kurz-, mittel- und langfristig)',
        'Priorisierung: Abgleich der finanziellen Tragweite mit den im '
        'zentralen Risikomanagementsystem erfassten Unternehmensrisiken',
        'Überwachung: Regelmäßige Aktualisierung im Rahmen des IKS und '
        'Integration in das Risikomanagement',
    ])

# ═══════════════════════════════════════════════════════════════════════════
# IRO-1 — [151] Beispiele finanzielle Risiken/Chancen → Aufzählung
# ═══════════════════════════════════════════════════════════════════════════

# p151 contains BOTH the intro "Das Verfahren berücksichtigt..." AND the examples
# → keep the intro text, strip inline list, add proper list items after
p151 = find_para('Das Verfahren berücksichtigt sowohl externe Abhängigkeiten')
if p151:
    set_text(p151,
        'Das Verfahren berücksichtigt sowohl externe Abhängigkeiten als auch '
        'die sich daraus ergebenden finanziellen Risiken und Chancen '
        '(vgl.\u202fESRS 1, Rz.\u202f78–79). Im Einzelnen wurden folgende '
        'Felder betrachtet:'
    )
    add_list_items_after(p151, [
        'Arbeitsmarkt und Demografie: Risiko des Fachkräftemangels mit '
        'finanziellen Auswirkungen auf Betriebssicherheit und Auslastung; '
        'Chance durch Positionierung als attraktiver Arbeitgeber mit '
        'guten Arbeitsbedingungen',
        'Regulatorische Rahmenbedingungen: Risiko steigender Kosten durch '
        'neue Vorgaben in der frühkindlichen Bildung oder arbeitsrechtliche '
        'Veränderungen; Chance durch staatliche Förderprogramme und Zuschüsse',
        'Energiepreise: Risiko deutlicher Kostensteigerungen für Strom und '
        'Wärme; Chance durch Investitionen in Effizienzmaßnahmen und '
        'langfristige Verträge mit erneuerbaren Energien',
        'Lebensmittelversorgung: Risiko höherer Einkaufspreise und '
        'Lieferengpässe; Chance durch nachhaltige Beschaffungsstrategien '
        'und regionale Lieferketten',
        'Gesellschaftliche Trends: Risiko sinkender Nachfrage bei '
        'rückläufigen Geburtenzahlen; Chance steigender Nachfrage durch '
        'den gesellschaftlichen Stellenwert früher Bildung',
    ])

# ═══════════════════════════════════════════════════════════════════════════
# IRO-1 — [156] Priorisierung → Fließtext
# ═══════════════════════════════════════════════════════════════════════════

p156 = find_para('eine Risikomatrix (Eintrittswahrscheinlichkeit')
if p156 and '- ' in p156.text:
    set_text(p156,
        'Die Priorisierung erfolgte über eine Risikomatrix '
        '(Eintrittswahrscheinlichkeit × Ausmaß), den Abgleich mit finanziellen '
        'Schwellenwerten (Budget, Liquidität und Wirtschaftlichkeit) sowie die '
        'Integration in die reguläre Risikoberichterstattung.'
    )

# ═══════════════════════════════════════════════════════════════════════════
# IRO-1 — [165]–[168] Datenquellen → Aufzählungen
# ═══════════════════════════════════════════════════════════════════════════

# [165] intro → keep; [166] interne Datenquellen → Aufzählung
# p166 contains "Interne Datenquellen: - Treibhausgasbilanz..." as ONE paragraph
# → modify text in-place, then add list items after it
p166 = find_para('Treibhausgasbilanz (Scope 1')
if p166 and '- ' in p166.text:
    set_text(p166, 'Als interne Datenquellen wurden herangezogen:')
    add_list_items_after(p166, [
        'Treibhausgasbilanz (Scope 1–3) einschließlich Energieverbrauch, '
        'Lebensmittel und Mobilität',
        'Personalstatistiken (z.\u202fB. Mitarbeiterzahlen, Fluktuation, '
        'Fachkräftebedarf)',
        'Nutzungs- und Belegungszahlen der Einrichtungen',
        'Ergebnisse aus Mitarbeiter- und Familienbefragungen',
        'Dokumentation im internen Kontrollsystem (IKS)',
    ])

# [167] externe Datenquellen → Fließtext
p167 = find_para('Demografische Kennzahlen')
if p167 and '- ' in p167.text:
    set_text(p167,
        'Als externe Datenquellen flossen demografische Kennzahlen '
        '(z.\u202fB. Geburtenraten, Bevölkerungsentwicklung), Entwicklungen '
        'der Energie- und Lebensmittelpreise sowie regulatorische '
        'Rahmenbedingungen (z.\u202fB. Vorgaben zur frühkindlichen Bildung, '
        'Arbeits- und Klimaschutzgesetze) in die Analyse ein.'
    )

# [168] "Externe Expertise: - ..." → Satz
p168 = find_para('Externe Expertise:')
if p168 and 'Planted' in p168.text:
    set_text(p168,
        'Die methodische Beratung und Validierung des Prozesses erfolgte durch '
        'die Nachhaltigkeitsagentur Planted.'
    )

# ═══════════════════════════════════════════════════════════════════════════
# SBM-1 — Quasi-Duplikate in Wertschöpfungskettenbeschreibung entfernen
# [107]/[113] (vorgelagerte WSK) und [106]/[114] (eigene Stufe) verschmelzen
# ═══════════════════════════════════════════════════════════════════════════

# "Fröbel ist als Träger von Kindertageseinrichtungen zentral positioniert.
#  Die vorgelagerte WSK umfasst insbesondere Lieferanten und Dienstleister für
#  Gebäude, Energie..."  → quasi-Duplikat von para "Vorgelagerte Wertschöpfung
#  umfasst insbesondere die Bereitstellung von Gebäuden und Infrastruktur..."
p_positionierung = find_para('Fröbel ist als Träger von Kindertageseinrichtungen zentral in der eigenen')
p_eigene_stufe2  = find_para('Die eigene Wertschöpfungsstufe besteht in der Organisation und Durchführung')
if p_positionierung:
    del_para(p_positionierung)
if p_eigene_stufe2:
    del_para(p_eigene_stufe2)

# ═══════════════════════════════════════════════════════════════════════════
# IRO-1 — kurze Einzel-Absätze "Eintrittswahrscheinlichkeit" / "Ausmaß"
# (standalone Normal-Absätze aus dem Finanzteil) in den Einleitungssatz
# integrieren — nur Normal-Absätze, nicht die List-Bullet-Scoring-Kriterien
# ═══════════════════════════════════════════════════════════════════════════

def find_para_normal(fragment):
    """Find first Normal-style paragraph containing fragment."""
    for p in doc.paragraphs:
        if fragment in p.text and p.style.name == 'Normal':
            return p
    return None

p_skala = find_para_normal('finanziellen Risiken und Chancen wurden anhand einer 5-Punkte-Skala')
p_ewk   = find_para_normal('Eintrittswahrscheinlichkeit: von sehr unwahrscheinlich')
p_ausm  = find_para_normal('Ausmaß: von minimal (1) bis sehr hoch')
if p_skala and p_ewk and p_ausm:
    set_text(p_skala,
        'Die finanziellen Risiken und Chancen wurden anhand einer 5-Punkte-Skala '
        'für Eintrittswahrscheinlichkeit und Ausmaß bewertet '
        '(vgl.\u202fESRS 1, Rz.\u202f82–83): '
        'Eintrittswahrscheinlichkeit von sehr unwahrscheinlich (1) bis '
        'sicherer Eintritt (5), Ausmaß von minimal (1) bis sehr hoch (5).'
    )
    del_para(p_ewk)
    del_para(p_ausm)

# ═══════════════════════════════════════════════════════════════════════════
# IRO-2 — Tabellenintro + Tippfehler
# ═══════════════════════════════════════════════════════════════════════════

p174 = find_para('IRO-2 Inhaltstabelle der Offenlegungsanforderungen')
if p174:
    set_text(p174,
        'Die folgende Tabelle gibt einen Überblick über die im Rahmen dieser '
        'Nachhaltigkeitserklärung abgedeckten Offenlegungsanforderungen.'
    )

# [177] Tippfehler
p177 = find_para('Klimwandel')
if p177:
    corrected = p177.text.replace('Klimwandel', 'Klimawandel').replace('wesentlch', 'wesentlich')
    set_text(p177, corrected)

# ═══════════════════════════════════════════════════════════════════════════
# Save
# ═══════════════════════════════════════════════════════════════════════════

doc.save(OUTPUT)
print(f"✓ Gespeichert: {OUTPUT}")
