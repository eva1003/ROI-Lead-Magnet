#!/usr/bin/env python3
"""Transform E1 (Umwelt) section of CSRD_Report_v7.docx – text quality improvements."""

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = 'Fröbel 2024/CSRD_Report_v7.docx'
OUTPUT = 'Fröbel 2024/CSRD_Report_v7.docx'

doc = Document(INPUT)

# ── Helpers ────────────────────────────────────────────────────────────────────

def find_para(fragment, skip=0):
    seen = 0
    for p in doc.paragraphs:
        if fragment in p.text:
            if seen == skip:
                return p
            seen += 1
    return None

def find_all(fragment):
    return [p for p in doc.paragraphs if fragment in p.text]

def del_para(para):
    e = para._element
    e.getparent().remove(e)

def del_para_and_next_empty(para):
    """Delete para and the immediately following empty paragraph (if any)."""
    e = para._element
    parent = e.getparent()
    nxt = e.getnext()
    parent.remove(e)
    if nxt is not None and nxt.tag.endswith('}p'):
        text_nodes = nxt.findall('.//' + qn('w:t'))
        if not ''.join(t.text or '' for t in text_nodes).strip():
            parent.remove(nxt)

def set_para_text(para, text):
    """Replace all runs in paragraph with a single run containing text."""
    p_elem = para._element
    for r in list(p_elem.findall(qn('w:r'))):
        p_elem.remove(r)
    for hl in list(p_elem.findall(qn('w:hyperlink'))):
        p_elem.remove(hl)
    new_r = OxmlElement('w:r')
    new_t = OxmlElement('w:t')
    new_t.text = text
    if text.startswith(' ') or text.endswith(' '):
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    new_r.append(new_t)
    p_elem.append(new_r)

LIST_STYLE = 'List Bullet'

def add_list_items_after(ref_para, items):
    """Insert List Bullet paragraphs after ref_para in order."""
    prev = ref_para
    for item in items:
        new_p = doc.add_paragraph()
        new_p.style = doc.styles[LIST_STYLE]
        new_p.add_run(item)
        prev._element.addnext(new_p._element)
        prev = new_p
    return prev

def strip_prefix(para, prefix):
    """Remove a prefix string from paragraph text."""
    t = para.text
    if t.startswith(prefix):
        set_para_text(para, t[len(prefix):])

def strip_qa_prefix(para):
    """Remove Q&A prefix 'Question?: Yes ' or 'Question?: No' from paragraph."""
    t = para.text
    for marker in [': Yes ', ': No ']:
        idx = t.find(marker)
        if idx != -1:
            set_para_text(para, t[idx + len(marker):])
            return
    # Fallback: 'Question?: No' at end of paragraph
    if t.endswith(': No'):
        del_para(para)

def split_inline_bullets(para, intro_text, items):
    """Set para to intro_text, then add List Bullet items after it."""
    set_para_text(para, intro_text)
    add_list_items_after(para, items)

# ── Fix 1: [208] Delete "trifft zu." standalone Boolean ───────────────────────
p = find_para('Erläuterung des Transitionsplans für den Klimaschutz trifft zu.')
if p and p.text.strip() == 'Erläuterung des Transitionsplans für den Klimaschutz trifft zu.':
    del_para(p)
    print("✓ [208] Deleted 'trifft zu.' Boolean")

# ── Fix 2: [214] Scope 1 inline bullets ───────────────────────────────────────
p214 = find_para('Scope 1 – Direkte Emissionen (Heizung & Fuhrpark) - Heizung')
if p214:
    t = p214.text
    # Split on " - Heizung " and " - Fuhrpark "
    part1 = t.split(' - Heizung ', 1)
    intro = part1[0].strip()  # "Scope 1 – Direkte Emissionen (Heizung & Fuhrpark)"
    rest  = 'Heizung ' + part1[1] if len(part1) > 1 else ''
    subitems = rest.split(' - Fuhrpark ', 1)
    item1 = subitems[0].strip()
    item2 = ('Fuhrpark ' + subitems[1].strip()) if len(subitems) > 1 else ''
    split_inline_bullets(p214, intro + ':', [item1, item2] if item2 else [item1])
    print("✓ [214] Scope 1 bullets fixed")

# ── Fix 3: [215] Scope 2 inline bullets ───────────────────────────────────────
p215 = find_para('Scope 2 – Indirekte Emissionen (Strom, Fernwärme) - Eingekaufte')
if p215:
    t = p215.text
    part1 = t.split(' - Eingekaufte Wärme', 1)
    intro = part1[0].strip()
    rest  = 'Eingekaufte Wärme' + part1[1] if len(part1) > 1 else ''
    subitems = rest.split(' - Eingekaufter Strom', 1)
    item1 = subitems[0].strip()
    item2 = ('Eingekaufter Strom' + subitems[1].strip()) if len(subitems) > 1 else ''
    split_inline_bullets(p215, intro + ':', [item1, item2] if item2 else [item1])
    print("✓ [215] Scope 2 bullets fixed")

# ── Fix 4: [216] Scope 3 inline bullets (6 items) ─────────────────────────────
p216 = find_para('Scope 3 – Weitere indirekte Emissionen - Kantine')
if p216:
    t = p216.text
    intro_end = t.find(' - Kantine')
    intro = t[:intro_end].strip()
    rest  = t[intro_end + 3:]  # skip ' - '
    # Split on " - " followed by known category starts
    categories = ['Pendeln (Mitarbeitende)', 'Drittfinanzierte', 'Geschäftsreisen', 'Abfall & Recycling', 'Sonstige Kategorien']
    items = [rest]
    for cat in categories:
        new_items = []
        for item in items:
            parts = item.split(f' - {cat}', 1)
            new_items.append(parts[0].strip())
            if len(parts) > 1:
                new_items.append(cat + parts[1])
        items = new_items
    items = [it.strip() for it in items if it.strip()]
    split_inline_bullets(p216, intro + ':', items)
    print(f"✓ [216] Scope 3 bullets fixed ({len(items)} items)")

# ── Fix 5: [217] Remove "Gesamteinschätzung " prefix ─────────────────────────
p217 = find_para('Gesamteinschätzung In Summe')
if p217:
    strip_prefix(p217, 'Gesamteinschätzung ')
    print("✓ [217] Removed 'Gesamteinschätzung' prefix")

# ── Fix 6: [219] Missing "wurde" + inline bullets ─────────────────────────────
p219 = find_para('Für das Jahr 2024 ein Globalbudget in Höhe von 1 Mio. Euro beschlossen')
if p219:
    t = p219.text
    # Fix grammar: add "wurde"
    t = t.replace('Für das Jahr 2024 ein Globalbudget', 'Für das Jahr 2024 wurde ein Globalbudget')
    # Split at "insbesondere: - "
    parts = t.split(': - der ', 1)
    if len(parts) == 2:
        intro = parts[0].strip() + ':'
        rest = 'der ' + parts[1]
        raw_items = [('der ' + x if not x.startswith('der ') and not x.startswith('dem ') and not x.startswith('sowie') else x)
                     for x in rest.split(', - ')]
        # Actually split on comma-dash pattern
        raw_items = [it.strip() for it in rest.split(' - ')]
        # Clean "sowie " prefix
        items = []
        for it in raw_items:
            items.append(it.strip().rstrip(',').rstrip('.'))
        items = [i for i in items if i]
        split_inline_bullets(p219, intro, items)
    else:
        set_para_text(p219, t)
    print("✓ [219] Fixed 'wurde' + bullets")

# ── Fix 7: [225] Em-dash inline list (Bestandsgebäude) ────────────────────────
p225 = find_para('Bestandsgebäuden mit konventionellen Heizsystemen (Gas, Öl, Fernwärme) – ')
if p225:
    t = p225.text
    # First token is category label, rest are em-dash items
    parts = [x.strip() for x in t.split(' – ')]
    cat_label = parts[0]  # "Bestandsgebäuden mit konventionellen Heizsystemen (Gas, Öl, Fernwärme)"
    items = parts[1:]
    intro = f'Ein wesentlicher Risikofaktor sind {cat_label.lower()}:'
    split_inline_bullets(p225, intro, items)
    print("✓ [225] Bestandsgebäude em-dash list fixed")

# ── Fix 8: [226] Em-dash inline list (Verpflegung) ────────────────────────────
p226 = find_para('Verpflegungssystem und Lebensmittelversorgung – ')
if p226:
    t = p226.text
    parts = [x.strip() for x in t.split(' – ')]
    cat_label = parts[0]
    items = parts[1:]
    intro = f'Auch das {cat_label} stellt einen strukturellen Risikofaktor dar:'
    split_inline_bullets(p226, intro, items)
    print("✓ [226] Verpflegung em-dash list fixed")

# ── Fix 9: [229] Q&A prefix ───────────────────────────────────────────────────
p229 = find_para('Ist das Unternehmen vom Paris- abgestimmten EU-Referenzwert ausgeschlossen?: Yes ')
if p229:
    strip_qa_prefix(p229)
    print("✓ [229] Q&A prefix removed")

# ── Fix 10: [231] Inline bullets after "insbesondere:" ────────────────────────
p231 = find_para('direkt mit den Handlungsfeldern der Geschäftsstrategie verknüpft – insbesondere: - Gebäude')
if p231:
    t = p231.text
    split_idx = t.find(' – insbesondere: - ')
    if split_idx != -1:
        intro = t[:split_idx + len(' – insbesondere:')].strip()
        rest  = t[split_idx + len(' – insbesondere: - '):]
        raw = ('Gebäude & Energie: ' if not rest.startswith('Gebäude') else '') + rest
        items = [x.strip() for x in raw.split(' - ') if x.strip()]
        split_inline_bullets(p231, intro, items)
    print("✓ [231] Bullets after 'insbesondere:' fixed")

# ── Fix 11: [234] Q&A prefix + typo "dder" ────────────────────────────────────
p234 = find_para('Wurde dieser von Verwaltungs-, Leitungs- und Aufsichtsorganen genehmigt?: Yes ')
if p234:
    strip_qa_prefix(p234)
    t = p234.text
    t = t.replace('un dder', 'und der').replace(' dder ', ' der ')
    set_para_text(p234, t)
    print("✓ [234] Q&A prefix removed, 'dder' typo fixed")

# ── Fix 12: [236]-[239] Progress blocks with inline bullets ───────────────────
progress_blocks = [
    ('Erzielte Fortschritte im Überblick: Energie & Gebäude: - Vollständige Umstellung',
     'Im Bereich Energie und Gebäude erzielte Fröbel folgende Fortschritte:'),
    ('Ernährung: - Fröbel-Leitfaden',
     'Im Bereich Ernährung wurden folgende Fortschritte erzielt:'),
    ('Mobilität: - Umsetzung eines nachhaltigen Mobilitätskonzepts',
     'Im Bereich Mobilität erzielte Fröbel folgende Ergebnisse:'),
    ('Monitoring & Steuerung: - Treibhausgasbilanzen',
     'Im Bereich Monitoring und Steuerung wurden folgende Maßnahmen umgesetzt:'),
]
for fragment, intro in progress_blocks:
    p = find_para(fragment)
    if p:
        t = p.text
        colon_idx = t.find(': - ')
        if colon_idx != -1:
            rest = t[colon_idx + 4:]  # skip ': - '
            items = [x.strip() for x in rest.split(' - ') if x.strip()]
            split_inline_bullets(p, intro, items)
            print(f"✓ Progress block fixed: {intro[:50]}")

# ── Fix 13: [240] Remove "Ausblick: " prefix ──────────────────────────────────
p240 = find_para('Ausblick: Für 2025–2026 liegt der Fokus')
if p240:
    strip_prefix(p240, 'Ausblick: ')
    print("✓ [240] 'Ausblick:' prefix removed")

# ── Fix 14: [244] Remove "trifft zu." + split policy list ─────────────────────
p244 = find_para('Richtlinien für Klimawandel trifft zu. Fröbel verfügt')
if p244:
    t = p244.text
    # Remove "Richtlinien für Klimawandel trifft zu. " sentence
    t_clean = t.replace('Richtlinien für Klimawandel trifft zu. ', '', 1)
    # Split on "\n\n- " to extract intro and list items
    if '\n' in t_clean:
        parts = t_clean.split('\n')
        parts = [x.strip() for x in parts if x.strip()]
        # First part is intro, rest are "- Policy: desc"
        intro = parts[0]
        items = []
        for part in parts[1:]:
            if part.startswith('- '):
                items.append(part[2:].strip())
            elif part:
                items.append(part)
        set_para_text(p244, intro)
        if items:
            add_list_items_after(p244, items)
    else:
        set_para_text(p244, t_clean.strip())
    print(f"✓ [244] 'trifft zu.' removed, policy list split ({len(items) if 'items' in dir() else 0} items)")

# ── Fix 15: [246] Lowercase table intro ───────────────────────────────────────
p246 = find_para('berücksichtigung von klimaschutz')
if p246:
    set_para_text(p246, 'Die folgende Tabelle zeigt die Angaben zur Berücksichtigung von Klimaschutz, Anpassung an den Klimawandel, Energieeffizienz, erneuerbaren Energien und sonstigen Bereichen in den Unternehmensrichtlinien.')
    print("✓ [246] Table intro capitalized")

# ── Fix 16: [251] Remove "trifft zu." sentence ────────────────────────────────
p251 = find_para('Maßnahmen zum Klimawandel trifft zu.')
if p251:
    t = p251.text.replace('Maßnahmen zum Klimawandel trifft zu. ', '', 1)
    set_para_text(p251, t.strip())
    print("✓ [251] 'trifft zu.' removed")

# ── Fix 17: [254-262] Duplicate lowercase table intros (5×) ───────────────────
all_thg_intros = find_all('thg-emissionsreduktionsziele und basisdaten')
if all_thg_intros:
    # Fix the first one
    set_para_text(all_thg_intros[0],
                  'Die folgende Tabelle zeigt die Angaben zu den THG-Emissionsreduktionszielen und Basisdaten.')
    # Delete remaining duplicates (+ their trailing empty paragraphs)
    for dup in all_thg_intros[1:]:
        del_para_and_next_empty(dup)
    print(f"✓ [254-262] THG table intros: kept 1, deleted {len(all_thg_intros)-1} duplicates")

# ── Fix 18: [265] Inline bullet list after "sichergestellt:" ──────────────────
p265 = find_para('Kohärenz zwischen Zielsystem und Inventar wird durch folgende Mechanismen sichergestellt: - Einheitliche')
if p265:
    t = p265.text
    split_idx = t.find(': - Einheitliche')
    if split_idx != -1:
        intro = t[:split_idx + 1].strip()  # keep up to ":"
        rest  = t[split_idx + 3:]          # skip ': - '
        # Find the concluding sentence " - Damit wird sichergestellt"
        damit_idx = rest.find(' - Damit wird sichergestellt')
        if damit_idx != -1:
            items_text = rest[:damit_idx]
            conclusion = 'Damit wird sichergestellt' + rest[damit_idx + len(' - Damit wird sichergestellt'):]
        else:
            items_text = rest
            conclusion = None
        items = [x.strip() for x in items_text.split(' - ') if x.strip()]
        set_para_text(p265, intro)
        last = add_list_items_after(p265, items)
        if conclusion:
            new_p = doc.add_paragraph()
            new_p.style = doc.styles['Normal']
            new_p.add_run(conclusion.strip())
            last._element.addnext(new_p._element)
    print("✓ [265] Coherence bullets fixed")

# ── Fix 19: [266] Delete "THG-Emissionsreduktionsziele trifft zu." ─────────────
p266 = find_para('THG-Emissionsreduktionsziele und Basisdaten trifft zu.')
if p266 and 'trifft zu.' in p266.text and len(p266.text.strip()) < 60:
    del_para(p266)
    print("✓ [266] Deleted 'trifft zu.' standalone")

# ── Fix 20: [268-271] Numbered action blocks with inline bullets ───────────────
action_blocks = [
    ('1. Energie & Gebäude (Scope 1 + 2) - Energieeffizienz',
     '1. Energie und Gebäude (Scope 1 + 2):'),
    ('2. Ernährung & Beschaffung (Scope 3) - Nachhaltige Verpflegung',
     '2. Ernährung und Beschaffung (Scope 3):'),
    ('3. Mobilität (Scope 1 + 3) - Elektrifizierung',
     '3. Mobilität (Scope 1 + 3):'),
    ('4. Monitoring & Steuerung - Jährliche THG-Bilanzierung',
     '4. Monitoring und Steuerung:'),
]
for fragment, new_intro in action_blocks:
    p = find_para(fragment)
    if p:
        t = p.text
        # Find position of first " - " that starts the bullet list
        dash_idx = t.find(' - ', t.find(') ') if ') ' in t else 0)
        if dash_idx == -1:
            dash_idx = t.find(': - ')
            if dash_idx != -1:
                dash_idx += 2  # point to " - "
        if dash_idx != -1:
            rest = t[dash_idx + 3:]  # skip " - "
            items = [x.strip() for x in rest.split(' - ') if x.strip()]
            split_inline_bullets(p, new_intro, items)
            print(f"✓ Action block fixed: {new_intro}")

# ── Fix 21: [279] Inline bullets after "würde Fröbel:" ────────────────────────
p279 = find_para('In einem solchen Fall würde Fröbel: - den neuen Bezugswert')
if p279:
    t = p279.text
    parts = t.split(': - ', 1)
    intro = parts[0].strip() + ':'
    items = [x.strip() for x in parts[1].split(' - ')] if len(parts) > 1 else []
    split_inline_bullets(p279, intro, items)
    print("✓ [279] 'würde Fröbel:' bullets fixed")

# ── Fix 22: [283] "Gesamtenergieverbrauch" → "Der Gesamtenergieverbrauch" ──────
p283 = find_para('Gesamtenergieverbrauch beläuft sich im Berichtszeitraum auf 21.635,9 MWh')
if p283:
    t = p283.text
    if not t.startswith('Der '):
        set_para_text(p283, 'Der ' + t)
    print("✓ [283] Added 'Der' to Gesamtenergieverbrauch")

# ── Fix 23: [284,286,288] Duplicate lowercase table intros (3×) ───────────────
all_energy_intros = find_all('Die folgende Tabelle zeigt die Angaben zu gesamtenergieverbrauch')
if all_energy_intros:
    set_para_text(all_energy_intros[0], 'Die folgende Tabelle zeigt den Gesamtenergieverbrauch.')
    for dup in all_energy_intros[1:]:
        del_para_and_next_empty(dup)
    print(f"✓ [284-288] Gesamtenergieverbrauch table intros: kept 1, deleted {len(all_energy_intros)-1}")

# ── Fix 24: [290-292] Fix garbled opening sentences; strip duplicate methodology ─
METHOD_MARKER = 'Für die Berechnung der Energieverbräuche wendet Fröbel ein abgestuftes Vorgehen'

def strip_method_appendix(para):
    """Remove everything from METHOD_MARKER onward in paragraph text."""
    t = para.text
    idx = t.find(METHOD_MARKER)
    if idx != -1:
        cleaned = t[:idx].strip()
        if cleaned:
            set_para_text(para, cleaned)
        else:
            del_para(para)
        return True
    return False

p290 = find_para('Für gesamtenergieverbrauch werden im Berichtsjahr 26,3 MWh ausgewiesen')
if p290:
    set_para_text(p290, 'Der Anteil erneuerbarer Energien am Gesamtenergieverbrauch beträgt im Berichtsjahr 26,3 MWh.')
    print("✓ [290] Garbled sentence fixed")

p291 = find_para('Gesamtenergieverbrauch beläuft sich im Berichtszeitraum auf 5.534,1 MWh')
if p291:
    # Remove methodology appendix (it's already in p283), fix label
    t = p291.text
    idx = t.find(METHOD_MARKER)
    intro = t[:idx].strip() if idx != -1 else t.strip()
    if not intro.startswith('Der '):
        intro = 'Der ' + intro
    set_para_text(p291, intro)
    print("✓ [291] Fixed Gesamtenergieverbrauch label")

p292 = find_para('Im Berichtsjahr beträgt gesamtenergieverbrauch 90,8 MWh')
if p292:
    t = p292.text
    idx = t.find(METHOD_MARKER)
    set_para_text(p292, 'Im Berichtsjahr beträgt der Anteil anderer erneuerbarer Energiequellen 90,8 MWh.')
    print("✓ [292] Fixed garbled 'gesamtenergieverbrauch' label")

# Delete [293] methodology note (already covered by p283)
p293 = find_para('Methodisch ist anzumerken: Für die Berechnung der Energieverbräuche wendet Fröbel')
if p293:
    del_para(p293)
    print("✓ [293] Deleted duplicate methodology note")

# ── Fix 25: Lowercase table intros ────────────────────────────────────────────
table_intro_fixes = [
    ('angabe der treibhausgasemissionen',
     'Die folgende Tabelle zeigt die Angaben zu den Treibhausgasemissionen.'),
    ('aufschlüsselung der scope-1- und scope-2-emissionen',
     'Die folgende Tabelle zeigt die Aufschlüsselung der Scope-1- und Scope-2-Emissionen.'),
    ('aufschlüsselung der treibhausgasemissionen gemäß esrs',
     'Die folgende Tabelle zeigt die Aufschlüsselung der Treibhausgasemissionen gemäß ESRS 1.'),
]
for fragment, fixed_text in table_intro_fixes:
    p = find_para(fragment)
    if p:
        set_para_text(p, fixed_text)
        print(f"✓ Table intro fixed: {fixed_text[:60]}")

# ── Fix 26: [309] Garbled Scope-3 opening sentence ────────────────────────────
p309 = find_para('berichtete Wert für zusammenstellung von Scope-3-THG-Bruttoemissionen beträgt 17')
if p309:
    t = p309.text
    # Replace the garbled opening sentence
    old_sent = 'Der berichtete Wert für zusammenstellung von Scope-3-THG-Bruttoemissionen beträgt 17.'
    new_sent = 'Die Primärdatenquote der Scope-3-THG-Bruttoemissionen beträgt im Berichtsjahr 17 %.'
    set_para_text(p309, t.replace(old_sent, new_sent, 1))
    print("✓ [309] Scope-3 garbled sentence fixed")

# ── Fix 27: [310,314] Duplicate scope-3 table intros ─────────────────────────
scope3_intros = find_all('zusammenstellung von scope-3-thg-bruttoemissionen')
if scope3_intros:
    set_para_text(scope3_intros[0], 'Die folgende Tabelle zeigt die Zusammenstellung der Scope-3-THG-Bruttoemissionen.')
    for dup in scope3_intros[1:]:
        del_para_and_next_empty(dup)
    print(f"✓ Scope-3 table intros: kept 1, deleted {len(scope3_intros)-1} duplicates")

# ── Fix 28: [316] Remove "Ergänzend ist festzuhalten: " prefix ────────────────
p316 = find_para('Ergänzend ist festzuhalten: Fröbel berichtet Scope-3 gemäß GHG-Protocol')
if p316:
    strip_prefix(p316, 'Ergänzend ist festzuhalten: ')
    print("✓ [316] 'Ergänzend ist festzuhalten:' prefix removed")

# ── Fix 29: [318] Garbled Scope-1 sentence ────────────────────────────────────
p318 = find_para('Angabe der scope-1-thg-bruttoemissionen beläuft sich')
if p318:
    set_para_text(p318,
        'Die Scope-1-THG-Bruttoemissionen belaufen sich im Berichtsjahr auf 1.331,3 t CO₂e; '
        'davon entfallen 9,8 t CO₂e auf Emissionen aus Emissionshandelssystemen.')
    print("✓ [318] Scope-1 garbled sentence fixed")

# ── Fix 30: [320] Garbled Scope-2 Q&A sentence ────────────────────────────────
p320 = find_para('Was sind ihre standortbezogenen scope-2-thg-bruttoemissionen')
if p320:
    set_para_text(p320,
        'Die standortbezogenen Scope-2-THG-Bruttoemissionen belaufen sich im Berichtsjahr auf 4.060,8 t CO₂e. '
        'Die standortbezogenen THG-Gesamtemissionen betragen 15.299,9 t CO₂e, '
        'die marktbezogenen Scope-2-THG-Bruttoemissionen 2.923,2 t CO₂e '
        'und die marktbezogenen THG-Gesamtemissionen 13.973,2 t CO₂e.')
    print("✓ [320] Scope-2 garbled Q&A sentence fixed")

# ── Fix 31: [327] Garbled THG-Gesamtemissionen sentence ──────────────────────
p327 = find_para('Was sind ihre standortbezogenen thg-gesamtemissionen')
if p327:
    set_para_text(p327,
        'Die standortbezogenen THG-Gesamtemissionen belaufen sich im Berichtsjahr auf 15.299,9 t CO₂e; '
        'die marktbezogenen THG-Gesamtemissionen betragen 13.973,2 t CO₂e.')
    print("✓ [327] THG-Gesamtemissionen garbled sentence fixed")

# ── Fix 32: [342,343,344] AR 45 Scope-2 instrument paragraphs ────────────────
p342 = find_para('Informationen zu Scope-2-THG-Bruttoemissionen und Energieinstrumenten beläuft sich im Berichtszeitraum auf 21,9')
if p342:
    t = p342.text
    # Fix garbled opening, keep methodology content after it
    method_start = t.find('Methodik & Anteile')
    if method_start != -1:
        method_rest = t[method_start:]
        set_para_text(p342, 'Der Anteil der durch vertragliche Instrumente (Herkunftsnachweise) gedeckten Scope-2-THG-Emissionen beträgt im Berichtszeitraum 21,9 %. ' + method_rest)
    print("✓ [342] AR 45 opening sentence fixed")

p343 = find_para('Informationen zu Scope-2-THG-Bruttoemissionen und Energieinstrumenten beläuft sich im Berichtszeitraum auf 100')
if p343:
    t = p343.text
    method_start = t.find('Methodik & Anteile')
    if method_start != -1:
        method_rest = t[method_start:]
        set_para_text(p343, 'Der durch vertragliche Instrumente gedeckte Anteil des Scope-2-Energieverbrauchs beträgt 100 %. ' + method_rest)
    print("✓ [343] AR 45 (100%) opening sentence fixed")

# Delete [344] – exact duplicate of [342]
p344 = find_para('Informationen zu Scope-2-THG-Bruttoemissionen und Energieinstrumenten beläuft sich im Berichtszeitraum auf 21,9', skip=1)
if p344 and 'Methodik & Anteile' in p344.text:
    del_para(p344)
    print("✓ [344] Deleted duplicate AR 45 paragraph")

# Delete [346],[347] – duplicated content from [342]
p346 = find_para('Der Anteil vertraglicher Instrumente (Ökostrom mit Herkunftsnachweisen, HKN) am gesamten Scope-2-Energieverbrauch ergibt sich aus')
if p346 and '[[REVIEW: explanation-driven]]' in p346.text:
    del_para(p346)
    print("✓ [346] Deleted duplicate methodology paragraph")

p347 = find_para('Ungebündelte Attribute werden nicht genutzt (0%). Fernwärme wird location-based', skip=0)
if p347 and '[[REVIEW: explanation-driven]]' in p347.text:
    del_para(p347)
    print("✓ [347] Deleted duplicate paragraph")

# Delete [349],[350],[351] – further duplicates
p349 = find_para('Methodik & Anteile (AR 45d/e): Fröbel berechnet Scope-2-Emissionen gemäß GHG-Protocol Scope-2-Guidance')
if p349 and '[[REVIEW: explanation-driven]]' in p349.text:
    del_para(p349)
    print("✓ [349] Deleted duplicate methodology paragraph")

p350 = find_para('Der Anteil vertraglicher Instrumente (Ökostrom mit Herkunftsnachweisen, HKN) am gesamten Scope-2-Energieverbrauch ergibt sich aus', skip=0)
if p350 and '[[REVIEW: explanation-driven]]' in p350.text:
    del_para(p350)
    print("✓ [350] Deleted duplicate paragraph")

p351 = find_para('Ungebündelte Attribute werden nicht genutzt (0%). Fernwärme wird location-based', skip=0)
if p351 and '[[REVIEW: explanation-driven]]' in p351.text:
    del_para(p351)
    print("✓ [351] Deleted duplicate paragraph")

# ── Fix 33: [358] Delete Q&A "No" Boolean (covered by [359]) ─────────────────
p358 = find_para('Hat Ihr Unternehmen Aktivitäten zum Abbau und zur Speicherung von Treibhausgasen?: No')
if p358 and p358.text.strip().endswith(': No'):
    del_para(p358)
    print("✓ [358] Deleted Q&A 'No' Boolean")

# ── Fix 34: [362] Lowercase table intro ───────────────────────────────────────
p362 = find_para('angaben zur entnahme und speicherung von treibhausgasen')
if p362:
    set_para_text(p362, 'Die folgende Tabelle zeigt die Angaben zur Entnahme und Speicherung von Treibhausgasen.')
    print("✓ [362] Table intro capitalized")

# ── Fix 35: [367] Lowercase table intro ───────────────────────────────────────
p367 = find_para('zusammenstellung der treibhausgasentnahme und -speicherung')
if p367:
    set_para_text(p367, 'Die folgende Tabelle zeigt die Zusammenstellung der Treibhausgasentnahme und -speicherung.')
    print("✓ [367] Table intro capitalized")

# ── Fix 36: [369] Q&A "No" Boolean – delete ───────────────────────────────────
p369 = find_para('Kommt es innerhalb der Unternehmenstätigkeiten und der vor- und nachgelagerten Wertschöpfungskette zur Entnahme')
if p369 and p369.text.strip().endswith(': No'):
    del_para(p369)
    print("✓ [369] Deleted Q&A 'No' Boolean")

# ── Fix 37: [372] Garbled CO₂-Zertifikate sentence ───────────────────────────
p372 = find_para('wie viele co2-zertifikate außerhalb der wertschöpfungskette wurden gelöscht')
if p372:
    set_para_text(p372,
        'Im Berichtsjahr wurden keine CO₂-Zertifikate außerhalb der Wertschöpfungskette gelöscht (0 t CO₂e); '
        'auch für die geplante Löschung von CO₂-Zertifikaten außerhalb der Wertschöpfungskette werden 0 t CO₂e ausgewiesen.')
    print("✓ [372] CO₂-Zertifikate garbled sentence fixed")

# ── Fix 38: [374-379] Delete 5 duplicate CO₂-Zertifikate paragraphs ──────────
cert_dups = find_all('Fröbel hat im Berichtszeitraum keine CO₂-Zertifikate erworben oder genutzt.')
if len(cert_dups) > 1:
    for dup in cert_dups[1:]:
        del_para(dup)
    print(f"✓ [375-379] Deleted {len(cert_dups)-1} duplicate CO₂-Zertifikate paragraphs")

# ── Fix 39: [382] Delete Q&A "No" (covered by [383]) ─────────────────────────
p382 = find_para('Hat das Unternehmen einen internen CO2-Preis?: No')
if p382 and p382.text.strip().endswith(': No'):
    del_para(p382)
    print("✓ [382] Deleted Q&A 'No' Boolean")

# ── Fix 40: [383] Remove "Zur Einordnung ist hinzuzufügen: " prefix ──────────
p383 = find_para('Zur Einordnung ist hinzuzufügen: Fröbel wendet zur Zeit keine internen CO2-Bepreisungssysteme')
if p383:
    strip_prefix(p383, 'Zur Einordnung ist hinzuzufügen: ')
    print("✓ [383] 'Zur Einordnung...' prefix removed")

# ── Fix 41: [385-387] Delete Q&A "No" paragraphs (covered by [388]) ──────────
qa_no_frags = [
    'Stimmen die internen CO2-Preise aus der Bewertung der Nutzungsdauer',
    'Stimmen die internen CO2-Preise aus der Wertminderung von Vermögenswerten',
    'Stimmen die internen CO2-Preise der Bemessung des beizulegenden Zeitwerts',
]
for frag in qa_no_frags:
    p = find_para(frag)
    if p and p.text.strip().endswith(': No'):
        del_para(p)
        print(f"✓ Deleted Q&A 'No': {frag[:50]}")

# ── Fix 42: [388] Remove "Zur Einordnung ist hinzuzufügen: " prefix ──────────
p388 = find_para('Zur Einordnung ist hinzuzufügen: Fröbel verwendet derzeit kein internes CO₂-Bepreisungssystem')
if p388:
    strip_prefix(p388, 'Zur Einordnung ist hinzuzufügen: ')
    print("✓ [388] 'Zur Einordnung...' prefix removed")

# ── Save ───────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"\n✓ Gespeichert: {OUTPUT}")
